"""T.C.A.M.S. — Targeted Contract Asymmetry Monitoring Summary.

Corpus-level DOCX report that aggregates the C.O.N.T.R.A. pilot corpus into
population-level findings across five dimensions:

  I.   CASI Score Distribution (band histogram)
  II.  Top-10 Entities by Aggregate CASI Score
  III. Axis Breakdown (corpus-wide axis dominance)
  IV.  Doctrinal Anchor Frequency (cross-entity prevalence)
  V.   L-Detector Heatmap (sub-detector firing frequency)

Requires an active SQLAlchemy Session connected to a DB containing
commercial_entities, commercial_documents, casi_scores, and contra_findings
tables (populated by Phase G ingest pipeline).

Style: Garamond, malachite (#1D6B44) headings, tan (#C8A882) sub-headings,
Page X of Y footer. No em-dash or en-dash in raw XML.

Source: C.O.N.T.R.A. Framework V1.0 Section VIII / Phase F Spec
"""

from __future__ import annotations

import logging
from datetime import UTC, datetime
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

if TYPE_CHECKING:
    from sqlalchemy.orm import Session

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Style constants (mirrors analytical_card.py palette)
# ---------------------------------------------------------------------------

_MALACHITE = RGBColor(0x1D, 0x6B, 0x44)
_TAN = RGBColor(0xC8, 0xA8, 0x82)
_FONT_BODY = "Garamond"
_FONT_FALLBACK = "Times New Roman"

# Band display order: low -> high adhesion
_BAND_ORDER = [
    "Baseline Adhesion",
    "Elevated Asymmetry",
    "Substantial Asymmetry",
    "Severe Asymmetry",
    "Foreclosure Regime",
]

_AXIS_LABELS = {
    "remedy_foreclosure": "Remedy Foreclosure",
    "data_extraction_depth": "Data Extraction Depth",
    "modification_and_consent": "Modification and Consent",
    "procedural_adhesion": "Procedural Adhesion",
    "enforcement_cost_asymmetry": "Enforcement Cost Asymmetry",
}


# ---------------------------------------------------------------------------
# Shared style helpers
# ---------------------------------------------------------------------------


def _run_font(run, size_pt: float = 11, bold: bool = False) -> None:
    run.font.name = _FONT_BODY
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), _FONT_BODY)
    rFonts.set(qn("w:hAnsi"), _FONT_BODY)
    rFonts.set(qn("w:cs"), _FONT_FALLBACK)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _section_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    _run_font(run, size_pt=14, bold=True)
    run.font.color.rgb = _MALACHITE
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)


def _sub_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _run_font(run, size_pt=11, bold=True)
    run.font.color.rgb = _TAN


def _body(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    for run in para.runs:
        _run_font(run)


def _kv(doc: Document, key: str, value: str) -> None:
    para = doc.add_paragraph()
    kr = para.add_run(f"{key}: ")
    _run_font(kr, bold=True)
    vr = para.add_run(value)
    _run_font(vr)


def _page_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run("Page ")
    _run_font(run, size_pt=9)
    for field_text in (" PAGE ", " NUMPAGES "):
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run._r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = field_text
        run._r.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)
        if field_text == " PAGE ":
            sep = para.add_run(" of ")
            _run_font(sep, size_pt=9)
            run = sep


# ---------------------------------------------------------------------------
# DB queries (pure functions taking session rows)
# ---------------------------------------------------------------------------


def _query_band_distribution(session: Session) -> dict[str, int]:
    from sqlalchemy import func

    from ..db.models import CasiScore

    rows = (
        session.query(CasiScore.band, func.count(CasiScore.document_hash))
        .group_by(CasiScore.band)
        .all()
    )
    dist = {b: 0 for b in _BAND_ORDER}
    for band, count in rows:
        if band in dist:
            dist[band] = count
        else:
            dist[band] = count
    return dist


def _query_top_entities(session: Session, limit: int = 10) -> list[tuple]:
    from sqlalchemy import func

    from ..db.models import CasiScore, CommercialDocument, CommercialEntity

    return (
        session.query(
            CommercialEntity.canonical_name,
            func.max(CasiScore.aggregate).label("max_score"),
            func.avg(CasiScore.aggregate).label("avg_score"),
            func.count(CasiScore.document_hash).label("doc_count"),
            CasiScore.band,
        )
        .join(
            CommercialDocument,
            CommercialDocument.entity_id == CommercialEntity.entity_id,
        )
        .join(CasiScore, CasiScore.document_hash == CommercialDocument.document_hash)
        .group_by(CommercialEntity.entity_id, CommercialEntity.canonical_name)
        .order_by(func.max(CasiScore.aggregate).desc())
        .limit(limit)
        .all()
    )


def _query_axis_sums(session: Session) -> dict[str, int | float]:
    from sqlalchemy import func

    from ..db.models import CasiScore

    row = session.query(
        func.sum(CasiScore.remedy_foreclosure),
        func.sum(CasiScore.data_extraction_depth),
        func.sum(CasiScore.modification_and_consent),
        func.sum(CasiScore.procedural_adhesion),
        func.sum(CasiScore.enforcement_cost_asymmetry),
        func.count(CasiScore.document_hash),
    ).one()
    keys = list(_AXIS_LABELS.keys()) + ["doc_count"]
    return dict(zip(keys, [v or 0 for v in row], strict=False))


def _query_anchor_frequency(session: Session, limit: int = 20) -> list[tuple]:
    from sqlalchemy import func

    from ..db.models import ContraFinding

    return (
        session.query(
            ContraFinding.doctrinal_anchor,
            func.count(ContraFinding.document_hash.distinct()).label("doc_count"),
            func.count(ContraFinding.id).label("finding_count"),
        )
        .group_by(ContraFinding.doctrinal_anchor)
        .order_by(func.count(ContraFinding.document_hash.distinct()).desc())
        .limit(limit)
        .all()
    )


def _query_detector_heatmap(session: Session, limit: int = 30) -> list[tuple]:
    from sqlalchemy import func

    from ..db.models import ContraFinding

    return (
        session.query(
            ContraFinding.layer,
            ContraFinding.sub_detector,
            func.count(ContraFinding.id).label("fire_count"),
        )
        .group_by(ContraFinding.layer, ContraFinding.sub_detector)
        .order_by(func.count(ContraFinding.id).desc())
        .limit(limit)
        .all()
    )


# ---------------------------------------------------------------------------
# Report builder
# ---------------------------------------------------------------------------


def build_tcams_report(
    session: Session,
    output_dir: str | Path,
    generated_by: str = "ODIA C.O.N.T.R.A. T.C.A.M.S.",
    framework_version: str = "1.0",
) -> Path:
    """Generate a T.C.A.M.S. corpus-level DOCX report.

    Arguments:
        session        -- active SQLAlchemy session with C.O.N.T.R.A. tables
        output_dir     -- directory to write the report
        generated_by   -- source label in the report header
        framework_version -- C.O.N.T.R.A. framework version tag

    Returns the resolved Path to the written .docx file.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now(UTC).strftime("%Y%m%d_%H%M%S")
    filename = f"tcams_report_{ts}.docx"
    output_path = output_dir / filename

    doc = Document()

    # ------------------------------------------------------------------
    # Title block
    # ------------------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_para.add_run(
        "T.C.A.M.S. -- TARGETED CONTRACT ASYMMETRY MONITORING SUMMARY"
    )
    _run_font(tr, size_pt=16, bold=True)
    tr.font.color.rgb = _MALACHITE

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_para.add_run(
        f"C.O.N.T.R.A. Framework V{framework_version} -- Corpus Population Report"
    )
    _run_font(sr, size_pt=11)
    sr.font.color.rgb = _TAN

    doc.add_paragraph()
    _kv(doc, "Generated", datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC"))
    _kv(doc, "Source", generated_by)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # Section I: CASI distribution
    # ------------------------------------------------------------------
    _section_heading(doc, "I. CASI SCORE DISTRIBUTION")
    _body(
        doc,
        "Distribution of ingested commercial documents by Consumer Adhesion Severity "
        "Index (CASI) band. Band thresholds: Baseline (0-20), Elevated (21-40), "
        "Substantial (41-60), Severe (61-80), Foreclosure (81-100).",
    )

    band_dist = _query_band_distribution(session)
    total_docs = sum(band_dist.values())

    tbl1 = doc.add_table(rows=1, cols=3)
    tbl1.style = "Table Grid"
    for i, h in enumerate(["Band", "Document Count", "Share"]):
        r = tbl1.rows[0].cells[i].paragraphs[0].add_run(h)
        _run_font(r, bold=True)
        r.font.color.rgb = _MALACHITE
    for band in _BAND_ORDER:
        count = band_dist.get(band, 0)
        share = f"{count / total_docs:.1%}" if total_docs else "N/A"
        row = tbl1.add_row()
        _run_font(row.cells[0].paragraphs[0].add_run(band))
        _run_font(row.cells[1].paragraphs[0].add_run(str(count)))
        _run_font(row.cells[2].paragraphs[0].add_run(share))

    # Totals row
    tot_row = tbl1.add_row()
    tot_r = tot_row.cells[0].paragraphs[0].add_run("TOTAL")
    _run_font(tot_r, bold=True)
    _run_font(tot_row.cells[1].paragraphs[0].add_run(str(total_docs)), bold=True)
    _run_font(tot_row.cells[2].paragraphs[0].add_run("100%"), bold=True)

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # Section II: Top-10 entities
    # ------------------------------------------------------------------
    _section_heading(doc, "II. TOP-10 ENTITIES BY AGGREGATE CASI SCORE")
    top_ents = _query_top_entities(session)
    if top_ents:
        tbl2 = doc.add_table(rows=1, cols=5)
        tbl2.style = "Table Grid"
        for i, h in enumerate(["Entity", "Max CASI", "Avg CASI", "Docs", "Band (max)"]):
            r = tbl2.rows[0].cells[i].paragraphs[0].add_run(h)
            _run_font(r, bold=True)
            r.font.color.rgb = _MALACHITE
        for name, max_score, avg_score, doc_count, band in top_ents:
            row = tbl2.add_row()
            _run_font(row.cells[0].paragraphs[0].add_run(name or "Unknown"))
            _run_font(row.cells[1].paragraphs[0].add_run(str(int(max_score or 0))))
            _run_font(row.cells[2].paragraphs[0].add_run(f"{avg_score or 0:.1f}"))
            _run_font(row.cells[3].paragraphs[0].add_run(str(doc_count)))
            _run_font(row.cells[4].paragraphs[0].add_run(band or ""))
    else:
        _body(doc, "No entities with CASI scores in corpus.")

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # Section III: Axis breakdown
    # ------------------------------------------------------------------
    _section_heading(doc, "III. AXIS BREAKDOWN -- CORPUS-WIDE DOMINANCE")
    _body(
        doc,
        "Sum of axis contributions across all scored documents. The dominant axis "
        "identifies the systemic asymmetry vector most prevalent across the corpus.",
    )

    axis_data = _query_axis_sums(session)
    doc_count_total = int(axis_data.get("doc_count", 0))
    _kv(doc, "Total scored documents", str(doc_count_total))
    doc.add_paragraph()

    axis_rows = [(label, axis_data.get(key, 0)) for key, label in _AXIS_LABELS.items()]
    axis_rows_sorted = sorted(axis_rows, key=lambda x: x[1], reverse=True)
    axis_total = sum(v for _, v in axis_rows) or 1

    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    for i, h in enumerate(["Axis", "Corpus Sum", "Share of Total"]):
        r = tbl3.rows[0].cells[i].paragraphs[0].add_run(h)
        _run_font(r, bold=True)
        r.font.color.rgb = _MALACHITE
    for label, val in axis_rows_sorted:
        row = tbl3.add_row()
        _run_font(row.cells[0].paragraphs[0].add_run(label))
        _run_font(row.cells[1].paragraphs[0].add_run(str(int(val))))
        _run_font(row.cells[2].paragraphs[0].add_run(f"{val / axis_total:.1%}"))

    if axis_rows_sorted:
        doc.add_paragraph()
        dominant = axis_rows_sorted[0]
        _body(
            doc,
            f"Dominant axis: {dominant[0]} (corpus sum {int(dominant[1])}, "
            f"{dominant[1] / axis_total:.1%} of total). This axis represents the "
            "primary mechanism by which these contracts foreclose consumer remedy.",
        )

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # Section IV: Doctrinal anchor frequency
    # ------------------------------------------------------------------
    _section_heading(doc, "IV. DOCTRINAL ANCHOR FREQUENCY (CROSS-ENTITY PREVALENCE)")
    _body(
        doc,
        "Doctrinal anchors appearing across the most unique commercial entities. "
        "High cross-entity prevalence indicates industry-wide adoption of a specific "
        "asymmetric term -- the empirical basis for the contracts-as-private-legislation "
        "thesis.",
    )

    anchors = _query_anchor_frequency(session)
    if anchors:
        tbl4 = doc.add_table(rows=1, cols=3)
        tbl4.style = "Table Grid"
        for i, h in enumerate(["Doctrinal Anchor", "Entities", "Findings"]):
            r = tbl4.rows[0].cells[i].paragraphs[0].add_run(h)
            _run_font(r, bold=True)
            r.font.color.rgb = _MALACHITE
        for anchor, doc_ct, finding_ct in anchors:
            row = tbl4.add_row()
            short = (
                (anchor[:80] + "...") if anchor and len(anchor) > 80 else (anchor or "")
            )
            _run_font(row.cells[0].paragraphs[0].add_run(short))
            _run_font(row.cells[1].paragraphs[0].add_run(str(doc_ct)))
            _run_font(row.cells[2].paragraphs[0].add_run(str(finding_ct)))
    else:
        _body(doc, "No findings with doctrinal anchors in corpus.")

    doc.add_paragraph()

    # ------------------------------------------------------------------
    # Section V: L-detector heatmap
    # ------------------------------------------------------------------
    _section_heading(doc, "V. L-DETECTOR HEATMAP -- SUB-DETECTOR FIRING FREQUENCY")
    _body(
        doc,
        "Sub-detectors ranked by total finding count. High-frequency sub-detectors "
        "identify the specific clause mechanisms most commonly deployed. Each "
        "sub-detector corresponds to a distinct doctrinal pattern in the "
        "C.O.N.T.R.A. controlled vocabulary.",
    )

    heatmap = _query_detector_heatmap(session)
    if heatmap:
        tbl5 = doc.add_table(rows=1, cols=3)
        tbl5.style = "Table Grid"
        for i, h in enumerate(["Layer", "Sub-Detector", "Firing Count"]):
            r = tbl5.rows[0].cells[i].paragraphs[0].add_run(h)
            _run_font(r, bold=True)
            r.font.color.rgb = _MALACHITE
        for layer, sub, count in heatmap:
            row = tbl5.add_row()
            _run_font(row.cells[0].paragraphs[0].add_run(layer or ""))
            _run_font(row.cells[1].paragraphs[0].add_run(sub or ""))
            _run_font(row.cells[2].paragraphs[0].add_run(str(count)))
    else:
        _body(doc, "No detector findings in corpus.")

    _page_footer(doc)
    doc.save(str(output_path))
    log.info("T.C.A.M.S. report written: %s", output_path)
    return output_path

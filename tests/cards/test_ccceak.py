"""Tests for cards.ccceak — C.C.C.E.A. clause exposure analysis report."""

import pytest

sqlalchemy = pytest.importorskip("sqlalchemy")
docx = pytest.importorskip("docx")

from datetime import UTC, datetime  # noqa: E402

from sqlalchemy import create_engine  # noqa: E402
from sqlalchemy.orm import sessionmaker  # noqa: E402

from oraculus_di_auditor.cards.ccceak import (  # noqa: E402
    _classify_clause,
    _cluster_excerpts,
    _jaccard,
    _probable_firm,
    _tokenize,
    build_ccceak_report,
)
from oraculus_di_auditor.db.models import (  # noqa: E402
    Base,
    CommercialDocument,
    CommercialEntity,
    ContraFinding,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def session(tmp_path):
    engine = create_engine(f"sqlite:///{tmp_path}/test_ccceak.db")
    Base.metadata.create_all(engine)
    SessionFactory = sessionmaker(bind=engine)  # noqa: N806
    sess = SessionFactory()
    yield sess
    sess.close()
    engine.dispose()


def _make_entity(entity_id: str, name: str) -> CommercialEntity:
    return CommercialEntity(
        entity_id=entity_id, canonical_name=name, in_contra_corpus=True
    )


def _make_doc(doc_hash: str, entity_id: str) -> CommercialDocument:
    return CommercialDocument(
        document_hash=doc_hash,
        entity_id=entity_id,
        doc_type="tos",
        retrieval_ts=datetime.now(UTC),
    )


def _make_finding(finding_id: str, doc_hash: str, excerpt: str) -> ContraFinding:
    return ContraFinding(
        finding_id=finding_id,
        document_hash=doc_hash,
        layer="L-11",
        sub_detector="A",
        severity="high",
        doctrinal_anchor="arbitration clause",
        evidence_excerpt=excerpt,
    )


@pytest.fixture()
def populated_session(session):
    e1 = _make_entity("eid-1", "Corp A")
    e2 = _make_entity("eid-2", "Corp B")
    session.add_all([e1, e2])

    d1 = _make_doc("aaa" + "0" * 61, "eid-1")
    d2 = _make_doc("bbb" + "0" * 61, "eid-2")
    session.add_all([d1, d2])

    # Near-identical arbitration excerpts (should cluster together)
    f1 = _make_finding(
        "f-001",
        d1.document_hash,
        "binding arbitration class action waiver final binding",
    )
    f2 = _make_finding(
        "f-002", d2.document_hash, "binding arbitration class action waiver final"
    )

    # Dissimilar excerpt (separate cluster)
    f3 = _make_finding(
        "f-003",
        d1.document_hash,
        "machine learning model training data purposes generative ai",
    )

    # Empty excerpt (should be excluded)
    f4 = _make_finding("f-004", d2.document_hash, "")

    session.add_all([f1, f2, f3, f4])
    session.commit()
    return session


# ---------------------------------------------------------------------------
# Unit: _tokenize
# ---------------------------------------------------------------------------


class TestTokenize:
    def test_basic(self):
        tokens = _tokenize("Binding Arbitration Clause")
        assert "binding" in tokens
        assert "arbitration" in tokens
        assert "clause" in tokens

    def test_strips_punctuation(self):
        tokens = _tokenize("class-action; waiver!")
        assert "class" in tokens or "classaction" in tokens
        assert "waiver" in tokens

    def test_empty_string(self):
        assert _tokenize("") == frozenset()


# ---------------------------------------------------------------------------
# Unit: _jaccard
# ---------------------------------------------------------------------------


class TestJaccard:
    def test_identical(self):
        s = frozenset(["a", "b", "c"])
        assert _jaccard(s, s) == pytest.approx(1.0)

    def test_disjoint(self):
        a = frozenset(["a", "b"])
        b = frozenset(["c", "d"])
        assert _jaccard(a, b) == pytest.approx(0.0)

    def test_partial_overlap(self):
        a = frozenset(["a", "b", "c"])
        b = frozenset(["b", "c", "d"])
        # intersection=2 union=4
        assert _jaccard(a, b) == pytest.approx(0.5)

    def test_both_empty(self):
        assert _jaccard(frozenset(), frozenset()) == pytest.approx(1.0)

    def test_one_empty(self):
        a = frozenset(["a"])
        assert _jaccard(a, frozenset()) == pytest.approx(0.0)


# ---------------------------------------------------------------------------
# Unit: _cluster_excerpts
# ---------------------------------------------------------------------------


class TestClusterExcerpts:
    def test_empty_input(self):
        assert _cluster_excerpts([]) == []

    def test_single_item(self):
        items = [("binding arbitration waiver", "L-11", "A", "f-001")]
        clusters = _cluster_excerpts(items)
        assert len(clusters) == 1
        assert len(clusters[0]) == 1

    def test_similar_items_cluster(self):
        items = [
            (
                "binding arbitration class action waiver final binding",
                "L-11",
                "A",
                "f-001",
            ),
            ("binding arbitration class action waiver final", "L-11", "A", "f-002"),
        ]
        clusters = _cluster_excerpts(items, threshold=0.35)
        assert len(clusters) == 1
        assert len(clusters[0]) == 2

    def test_dissimilar_items_separate(self):
        items = [
            ("binding arbitration class action waiver", "L-11", "A", "f-001"),
            (
                "machine learning training data generative ai model",
                "L-12",
                "B",
                "f-002",
            ),
        ]
        clusters = _cluster_excerpts(items, threshold=0.35)
        assert len(clusters) == 2

    def test_clusters_sorted_by_size_desc(self):
        items = [
            ("alpha beta gamma delta epsilon zeta", "L-11", "A", f"f-{i:03d}")
            for i in range(3)
        ] + [
            (
                "completely different text about machine learning model",
                "L-12",
                "B",
                "f-100",
            ),
        ]
        clusters = _cluster_excerpts(items, threshold=0.5)
        sizes = [len(c) for c in clusters]
        assert sizes == sorted(sizes, reverse=True)


# ---------------------------------------------------------------------------
# Unit: _classify_clause
# ---------------------------------------------------------------------------


class TestClassifyClause:
    def test_arbitration(self):
        result = _classify_clause("disputes resolved through binding arbitration only")
        assert result == "Mandatory Arbitration / Class Waiver"

    def test_ml_training(self):
        result = _classify_clause(
            "your data may be used to train models for machine learning"
        )
        assert result == "ML Training / AI Data Use"

    def test_data_collection(self):
        result = _classify_clause(
            "we collect data about your usage and aggregate data for analytics"
        )
        assert result == "Broad Data Collection / License"

    def test_unilateral_modification(self):
        result = _classify_clause(
            "we may amend any time without notice at our sole discretion"
        )
        assert result == "Unilateral Modification"

    def test_limitation_of_liability(self):
        result = _classify_clause(
            "our liability is limited and we disclaim all warranties"
        )
        assert result == "Limitation of Liability / Disclaimer"

    def test_unclassified(self):
        result = _classify_clause("ordinary sentence with no special keywords")
        assert result == "Unclassified Clause"


# ---------------------------------------------------------------------------
# Unit: _probable_firm
# ---------------------------------------------------------------------------


class TestProbableFirm:
    def test_no_match_returns_unknown(self):
        result = _probable_firm("ordinary contractual language")
        assert result == "Unknown / No Match"

    def test_cooley_pattern(self):
        result = _probable_firm(
            "bilateral arbitration mutual class waiver opt out thirty days"
        )
        assert result == "Cooley LLP"

    def test_latham_pattern(self):
        result = _probable_firm(
            "new york govern jams streamlined rules apply to disputes"
        )
        assert result == "Latham & Watkins LLP"


# ---------------------------------------------------------------------------
# Integration: build_ccceak_report
# ---------------------------------------------------------------------------


class TestBuildCccEakReport:
    def test_returns_path(self, populated_session, tmp_path):
        path = build_ccceak_report(populated_session, tmp_path)
        assert path.exists()
        assert path.suffix == ".docx"

    def test_creates_output_dir(self, populated_session, tmp_path):
        out = tmp_path / "reports" / "ccceak"
        path = build_ccceak_report(populated_session, out)
        assert out.exists()
        assert path.parent == out

    def test_empty_db_produces_valid_docx(self, session, tmp_path):
        path = build_ccceak_report(session, tmp_path)
        assert path.exists()
        assert path.stat().st_size > 0

    def test_empty_excerpt_excluded(self, populated_session, tmp_path):
        from docx import Document as DocxDoc

        path = build_ccceak_report(populated_session, tmp_path, jaccard_threshold=0.3)
        doc = DocxDoc(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        # 3 non-empty findings; empty f-004 excluded
        assert "Total Findings Loaded: 3" in text

    def test_filename_includes_ccceak(self, populated_session, tmp_path):
        path = build_ccceak_report(populated_session, tmp_path)
        assert "ccceak_report" in path.name

    def test_docx_contains_title_block(self, populated_session, tmp_path):
        from docx import Document as DocxDoc

        path = build_ccceak_report(populated_session, tmp_path)
        doc = DocxDoc(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "C.C.C.E.A." in text
        assert "D-13.H" in text

    def test_probable_firm_note_present(self, populated_session, tmp_path):
        from docx import Document as DocxDoc

        path = build_ccceak_report(populated_session, tmp_path)
        doc = DocxDoc(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "PROBABLE" in text

    def test_similar_excerpts_cluster(self, populated_session, tmp_path):
        from docx import Document as DocxDoc

        # Threshold 0.3 should cluster f-001 and f-002 (similar arb text)
        path = build_ccceak_report(populated_session, tmp_path, jaccard_threshold=0.3)
        doc = DocxDoc(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Cluster 1" in text
        # The largest cluster (arb pair) should report 2 findings
        assert "2 findings" in text

    def test_max_clusters_respected(self, populated_session, tmp_path):
        from docx import Document as DocxDoc

        path = build_ccceak_report(
            populated_session, tmp_path, max_clusters=1, jaccard_threshold=0.3
        )
        doc = DocxDoc(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Cluster 1" in text
        assert "Cluster 2" not in text

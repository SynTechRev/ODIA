"""FastAPI routes for file upload and audit pipeline execution.

Provides a complete web-based audit workflow:
  - Upload documents (PDF, JSON, TXT, XML)
  - Upload images with OCR extraction (JPEG, PNG)
  - List / delete uploaded files
  - Run the full audit pipeline in a background thread
  - Poll job progress and retrieve results
  - Download reports (Markdown, HTML, PDF, DOCX) and evidence packets (ZIP)

Endpoints:
  POST   /api/v1/upload                         upload single file
  POST   /api/v1/upload/batch                   upload multiple files
  POST   /api/v1/upload/image                   upload image and extract text via OCR
  GET    /api/v1/upload/files                   list uploaded files
  DELETE /api/v1/upload/files/{file_id}         remove a file
  POST   /api/v1/audit/run                      start audit job
  GET    /api/v1/audit/status/{job_id}          poll job progress
  GET    /api/v1/audit/results/{job_id}         retrieve job results
  GET    /api/v1/audit/export/{job_id}          download report
  GET    /api/v1/audit/evidence-packet/{job_id} download evidence ZIP
"""

from __future__ import annotations

import hashlib
import io
import json
import logging
import tempfile
import threading
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# In-memory stores — process-lifetime; suitable for local/single-user deployment
# ---------------------------------------------------------------------------
_FILES: dict[str, dict[str, Any]] = {}
_JOBS: dict[str, dict[str, Any]] = {}
_STORE_LOCK = threading.Lock()

# Temp directory — created once at import
_UPLOAD_DIR: Path = Path(tempfile.gettempdir()) / "odia_uploads"
_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

_ALLOWED_EXTENSIONS = frozenset(
    {".pdf", ".json", ".txt", ".xml", ".html", ".htm", ".doc", ".docx", ".tif", ".tiff"}
)
_ALLOWED_IMAGE_EXTENSIONS = frozenset({".jpg", ".jpeg", ".png"})

try:
    from fastapi import APIRouter, File, HTTPException, UploadFile
    from fastapi.responses import StreamingResponse
    from pydantic import BaseModel
    from pydantic import Field as PydanticField

    _FASTAPI_AVAILABLE = True
except ImportError:  # pragma: no cover
    _FASTAPI_AVAILABLE = False
    APIRouter = None  # type: ignore[assignment]
    BaseModel = object  # type: ignore[assignment, misc]
    PydanticField = lambda *a, **kw: None  # type: ignore[assignment]  # noqa: E731


# ---------------------------------------------------------------------------
# File ingestion helpers
# ---------------------------------------------------------------------------


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def register_uploaded_path(
    path: Path,
    *,
    source: str | None = None,
    move: bool = True,
) -> dict[str, Any]:
    """Register a file that already exists on disk into the upload store.

    Used by the Legistar retrieval flow (v2.7.6 X3) so that documents
    pulled from a city's Legistar portal land in the same ``_FILES``
    table as drag-and-drop uploads — and therefore appear in the Upload
    page's "files ready" list and are eligible for ``POST /audit/run``.

    Parameters
    ----------
    path:
        Existing file. Its name (with a fresh ``file_id`` prefix) is
        used as the destination filename inside ``_UPLOAD_DIR``.
    source:
        Optional free-form source identifier (e.g. the Legistar URL)
        recorded in the metadata for downstream provenance.
    move:
        When True (default) the file is moved into ``_UPLOAD_DIR``;
        when False it is copied. Move is the right default for
        Legistar staging dirs since the staging copy serves no
        further purpose after registration.

    Returns
    -------
    The same metadata dict the standard ``POST /upload`` route returns,
    so the frontend's ``listUploadedFiles()`` poll picks it up
    transparently.
    """
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"Cannot register missing file: {path}")

    ext = path.suffix.lower()
    if ext not in _ALLOWED_EXTENSIONS and ext not in _ALLOWED_IMAGE_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}' — Legistar retrieval should "
            "filter document types upstream."
        )

    file_id = str(uuid.uuid4())[:8]
    safe_name = path.name
    dest = _UPLOAD_DIR / f"{file_id}_{safe_name}"

    if move:
        path.rename(dest)
    else:
        dest.write_bytes(path.read_bytes())

    content = dest.read_bytes()
    meta: dict[str, Any] = {
        "file_id": file_id,
        "name": safe_name,
        "size": len(content),
        "sha256": _sha256_bytes(content),
        "format": ext.lstrip("."),
        "path": str(dest),
        "uploaded_at": datetime.now(UTC).isoformat(),
    }
    if source:
        meta["source"] = source

    with _STORE_LOCK:
        _FILES[file_id] = meta
    return meta


def ingest_uploaded_file(path: Path) -> dict[str, Any]:
    """Read an uploaded file and return a minimal document dict for analysis.

    Exported for testing.
    """
    ext = path.suffix.lower()
    text = ""
    text_extraction: dict[str, Any] | None = None

    if ext == ".txt":
        text = path.read_text(encoding="utf-8", errors="replace")

    elif ext == ".json":
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                text = (
                    data.get("raw_text")
                    or data.get("text")
                    or data.get("content")
                    or data.get("body")
                    or ""
                )
                if not text:
                    text = json.dumps(data, indent=2)
            else:
                text = json.dumps(data, indent=2)
        except Exception:
            text = path.read_text(encoding="utf-8", errors="replace")

    elif ext == ".xml":
        try:
            import xml.etree.ElementTree as ET  # noqa: N814

            tree = ET.parse(str(path))
            text = " ".join(
                (el.text or "").strip()
                for el in tree.iter()
                if el.text and el.text.strip()
            )
        except Exception:
            text = path.read_text(encoding="utf-8", errors="replace")

    elif ext == ".docx":
        # v3.2.5: modern Word (.docx) — pulled via python-docx, paragraphs
        # joined newline-delimited so detectors see the document as flowing
        # text. Tables flattened by row-then-cell ordering. Headers/footers
        # skipped since they are usually agency boilerplate that depresses
        # signal-to-noise without adding civic-procurement information.
        try:
            from docx import Document

            doc = Document(str(path))
            chunks: list[str] = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if t:
                    chunks.append(t)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        chunks.append(" | ".join(cells))
            text = "\n".join(chunks)
        except Exception:
            text = ""

    elif ext == ".doc":
        # v3.2.5: legacy Word binary (.doc / OLE compound format). No
        # reliable pure-Python parser exists, so we shell out to one of:
        #   1. antiword (Debian package: `apt-get install antiword`) —
        #      fastest, smallest dependency
        #   2. libreoffice --headless --convert-to txt — heavier but
        #      already on many servers
        # If neither is installed we fall back to empty text and log a
        # warning; the document still gets persisted (raw bytes + SHA-256
        # for provenance) but detectors see no content.
        import shutil
        import subprocess
        import tempfile

        text = ""
        if shutil.which("antiword"):
            try:
                r = subprocess.run(
                    ["antiword", "-w", "0", str(path)],
                    capture_output=True,
                    timeout=30,
                    check=False,
                )
                if r.returncode == 0:
                    text = r.stdout.decode("utf-8", errors="replace")
            except Exception:
                text = ""

        if not text:
            soffice = shutil.which("libreoffice") or shutil.which("soffice")
            if soffice:
                try:
                    with tempfile.TemporaryDirectory() as outdir:
                        subprocess.run(
                            [
                                soffice,
                                "--headless",
                                "--convert-to",
                                "txt",
                                "--outdir",
                                outdir,
                                str(path),
                            ],
                            capture_output=True,
                            timeout=60,
                            check=False,
                        )
                        out_txt = Path(outdir) / (path.stem + ".txt")
                        if out_txt.exists():
                            text = out_txt.read_text(encoding="utf-8", errors="replace")
                except Exception:
                    text = ""

        if not text:
            logger.warning(
                "ingest_uploaded_file: no .doc text extractor available "
                "(install antiword or libreoffice); doc %s will be stored "
                "with empty text",
                path.name,
            )

    elif ext in (".tif", ".tiff"):
        # v3.2.5: scanned multi-page TIFF (Tulare County's pre-2005 BOS
        # archive is mostly TIFFs — board minutes, budget hearings, addenda
        # were paper-scanned and microfilm-digitized). Iterate every frame
        # via PIL.Image.seek + OCR each via pytesseract, then concatenate
        # blank-line separated. Without seek() we'd only get the cover
        # page and miss everything after.
        try:
            import pytesseract  # type: ignore[import]
            from PIL import Image  # type: ignore[import]

            img = Image.open(str(path))
            pages: list[str] = []
            try:
                page_idx = 0
                while True:
                    img.seek(page_idx)
                    page_text = pytesseract.image_to_string(img)
                    if page_text and page_text.strip():
                        pages.append(page_text.strip())
                    page_idx += 1
            except EOFError:
                pass
            text = "\n\n".join(pages)
            text_extraction = {
                "method": "tesseract_ocr_tiff",
                "char_count": len(text),
                "ocr_used": True,
                "page_count": page_idx,
            }
        except ImportError:
            text = (
                f"[TIFF: {path.name} — install pytesseract + Pillow "
                f"for OCR text extraction]"
            )
            text_extraction = {
                "method": "pytesseract_unavailable",
                "char_count": 0,
                "ocr_used": False,
            }
        except Exception as exc:
            text = f"[TIFF: {path.name} — OCR error: {exc}]"
            text_extraction = {
                "method": "tesseract_ocr_tiff_failed",
                "char_count": 0,
                "ocr_used": False,
            }

    elif ext in (".html", ".htm"):
        # v3.1.1: HTML pages (e.g. WordPress press releases scraped via
        # /webhook/scrape-and-ingest-async). Strip tags + scripts + styles
        # via BeautifulSoup (already a dep) so detectors see body text
        # only, not navigation chrome or inline JS. Falls back to raw
        # read on parse failure so the audit always gets *some* signal.
        #
        # v3.2.4: prefer semantic content containers (<main>, <article>,
        # [role="main"]) when present. Drupal / Wagtail / modern CMS
        # themes wrap navigation in <div> elements with custom classes
        # that the v3.1.1 generic strip (script/style/nav/footer/aside)
        # misses entirely — observed live on Tulare County's Drupal
        # site where the article body is 1.6 KB of real prose buried
        # in 13 KB of <div>-wrapped sidebar/menu cruft. Semantic
        # extraction recovers just the article text. Falls back to
        # the generic strip if no semantic container is present (older
        # WordPress themes, hand-coded HTML).
        try:
            from bs4 import BeautifulSoup

            raw = path.read_text(encoding="utf-8", errors="replace")
            soup = BeautifulSoup(raw, "html.parser")

            # Strip scripts/styles globally so they don't leak into
            # whichever container we pick.
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()

            # Try semantic-container extraction first. ~200 chars is
            # the minimum that suggests we actually got an article
            # body (not just a placeholder div).
            text = ""
            for selector in ("main", "article", '[role="main"]'):
                el = soup.select_one(selector)
                if el is not None:
                    candidate = el.get_text(separator="\n", strip=True)
                    if len(candidate) >= 200:
                        text = candidate
                        break

            # No semantic container yielded enough — fall through to
            # the v3.1.1 generic strip (preserves WordPress / older
            # CMS behaviour).
            if not text:
                for tag in soup(["nav", "footer", "aside"]):
                    tag.decompose()
                text = soup.get_text(separator="\n", strip=True)
        except Exception:
            text = path.read_text(encoding="utf-8", errors="replace")

    if ext == ".pdf":
        # Delegate to ingestion.engine.extract_text_from_pdf_with_metadata
        # so scanned PDFs fall through to the Tesseract+Poppler OCR path
        # AND we record which extraction path actually fired (v2.9.3 A.3).
        try:
            from oraculus_di_auditor.ingestion.engine import (
                extract_text_from_pdf_with_metadata,
            )

            result = extract_text_from_pdf_with_metadata(path)
            text = result.text
            text_extraction = {
                "method": result.method,
                "char_count": result.char_count,
                "ocr_used": result.method == "tesseract_ocr",
            }
        except ImportError:
            text = f"[PDF: {path.name} — install pypdf to extract text]"
            text_extraction = {
                "method": "pypdf_unavailable",
                "char_count": 0,
                "ocr_used": False,
            }
        except Exception as exc:
            text = f"[PDF: {path.name} — extraction error: {exc}]"
            text_extraction = {
                "method": "failed",
                "char_count": 0,
                "ocr_used": False,
            }

    out: dict[str, Any] = {
        "document_id": path.stem,
        "raw_text": text,
        "title": path.name,
        "source": str(path),
    }
    if text_extraction is not None:
        out["text_extraction"] = text_extraction
    return out


def _ocr_image(path: Path) -> tuple[str, str]:
    """Extract text from an image file.

    Returns ``(raw_text, ocr_method)`` where *ocr_method* is one of
    ``"tesseract"``, ``"pillow_stub"``, or ``"unavailable"``.

    Exported for testing.
    """
    try:
        import pytesseract  # type: ignore[import]
        from PIL import Image  # type: ignore[import]

        img = Image.open(str(path))
        text = pytesseract.image_to_string(img)
        return text.strip(), "tesseract"
    except ImportError:
        pass
    except Exception as exc:
        logger.warning("tesseract OCR failed for %s: %s", path.name, exc)

    # Fallback: return filename as stub text so the document is still auditable
    return (
        f"[Image: {path.name} — install pytesseract for OCR text extraction]",
        "unavailable",
    )


def _flatten_findings(result: dict[str, Any], document_id: str) -> list[dict[str, Any]]:
    """Flatten the nested findings dict from analyze_document into a flat list."""
    flat: list[dict[str, Any]] = []
    raw = result.get("anomalies", result.get("findings", []))

    if isinstance(raw, dict):
        for detector, items in raw.items():
            if isinstance(items, list):
                for item in items:
                    if isinstance(item, dict):
                        f = dict(item)
                        f.setdefault("layer", detector)
                        f["document_id"] = document_id
                        flat.append(f)
    elif isinstance(raw, list):
        for item in raw:
            if isinstance(item, dict):
                f = dict(item)
                f["document_id"] = document_id
                flat.append(f)

    return flat


def _build_markdown_report(results: dict[str, Any]) -> str:
    """Build a Markdown audit report from job results."""
    doc_count = results.get("document_count", 0)
    finding_count = results.get("finding_count", 0)
    generated = results.get("generated_at", "N/A")
    sev = results.get("severity_summary", {})

    lines: list[str] = [
        "# O.D.I.A. Audit Report",
        "",
        f"**Generated**: {generated}  ",
        f"**Documents analyzed**: {doc_count}  ",
        f"**Total findings**: {finding_count}  ",
        "",
        "## Severity Summary",
        "",
    ]
    for level in ("critical", "high", "medium", "low"):
        count = sev.get(level, 0)
        if count:
            lines.append(f"- **{level.upper()}**: {count}")

    lines += ["", "---", "", "## Findings", ""]

    findings = sorted(
        results.get("findings", []),
        key=lambda f: {"critical": 0, "high": 1, "medium": 2, "low": 3}.get(
            f.get("severity", "low"), 3
        ),
    )

    for i, f in enumerate(findings, 1):
        severity = f.get("severity", "unknown").upper()
        lines += [
            f"### Finding {i:03d} — [{severity}] {f.get('issue', 'Unknown issue')}",
            "",
            f"**Detector**: `{f.get('layer', 'unknown')}`  ",
            f"**Document**: `{f.get('document_id', 'unknown')}`  ",
            f"**Severity**: {severity}  ",
            "",
        ]
        if "plain_summary" in f:
            lines += [
                f"**Summary**: {f['plain_summary']}  ",
                "",
                f"**Why it matters**: {f.get('plain_impact', '')}  ",
                "",
                f"**Recommended action**: {f.get('plain_action', '')}  ",
                "",
            ]
        lines += ["---", ""]

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Background audit executor
# ---------------------------------------------------------------------------


_MESH_AGENT_PIPELINE = (
    "ingestion",
    "analysis",
    "anomaly",
    "synthesis",
    "database",
    "interface",
)


def _record_mesh_job(job_id: str, file_count: int, status: str, **extras: Any) -> None:
    """Best-effort write of a MeshExecutionJob row for an audit run
    (v2.7.6 X4).

    The Orchestrator page's "Recent Mesh Jobs" panel reads from this
    table. Pre-X4 it stayed empty even after dozens of audits, because
    only the legacy n8n-coordinated path wrote to it. Recording the
    audit pipeline here makes the panel the single "what did ODIA
    just do?" surface for the desktop install. Failure to write is
    swallowed — the audit itself is the source of truth and must not
    be derailed by an observability side-effect.
    """
    try:
        from sqlalchemy.exc import IntegrityError

        from oraculus_di_auditor.db import models as db_models
        from oraculus_di_auditor.db.session import get_db
    except ImportError:
        return

    try:
        with get_db() as session:
            existing = (
                session.query(db_models.MeshExecutionJob)
                .filter(db_models.MeshExecutionJob.job_id == job_id)
                .one_or_none()
            )
            now = datetime.now(UTC).replace(tzinfo=None)
            if existing is None:
                row = db_models.MeshExecutionJob(
                    job_id=job_id,
                    job_type="audit",
                    status=status,
                    agent_count=len(_MESH_AGENT_PIPELINE),
                    task_count=file_count,
                    gcn_validated=False,
                    governor_approved=False,
                    started_at=now if status != "queued" else None,
                    completed_at=now if status in ("completed", "failed") else None,
                    metadata_json=json.dumps(
                        {
                            "pipeline": list(_MESH_AGENT_PIPELINE),
                            "source": "upload.audit_run",
                            **extras,
                        }
                    ),
                )
                session.add(row)
            else:
                existing.status = status
                if status in ("completed", "failed"):
                    existing.completed_at = now
                results = extras.pop("results", None)
                if results is not None:
                    existing.results_json = json.dumps(results)
                if extras:
                    try:
                        meta = json.loads(existing.metadata_json or "{}")
                    except Exception:
                        meta = {}
                    meta.update(extras)
                    existing.metadata_json = json.dumps(meta)
            session.commit()
    except IntegrityError:
        # Race condition — another worker beat us to the insert. Safe
        # to ignore; the row exists either way.
        return
    except Exception as exc:  # noqa: BLE001
        logger.warning("Failed to record MeshExecutionJob for %s: %s", job_id, exc)


def _persist_upload_document(
    doc: dict[str, Any],
    findings: list[dict[str, Any]],
    jurisdiction: str | None,
) -> None:
    """Persist an upload-flow document + its analysis to the DB.

    Best-effort — failure is logged but never propagates to the audit job.
    Skips if the document_id is already present (idempotent re-runs).
    """
    try:
        from oraculus_di_auditor.db import crud as db_crud
        from oraculus_di_auditor.db.session import get_db
    except ImportError:
        return

    try:
        with get_db() as session:
            existing = db_crud.get_document_by_id(session, doc["document_id"])
            if existing is None:
                db_crud.create_document(
                    session,
                    {
                        "document_id": doc["document_id"],
                        "title": doc.get("title") or doc.get("filename", "Untitled"),
                        "document_type": doc.get("document_type", "document"),
                        "jurisdiction": jurisdiction,
                        "authority": doc.get("authority"),
                        "version_date": doc.get("version_date"),
                        "signatory": doc.get("signatory"),
                    },
                )

            scalar = doc.get("scalar_score")
            db_crud.create_analysis(
                session,
                {
                    "document_id": doc["document_id"],
                    "scalar_score": float(scalar) if scalar is not None else 1.0,
                    "anomaly_count": len(findings),
                    "engine_version": "3.5.0",
                    "metadata": {"source": "upload.audit_run"},
                },
                [
                    {
                        "anomaly_id": f.get("id", "unknown"),
                        "issue": f.get("issue", ""),
                        "severity": f.get("severity", "low"),
                        "layer": f.get("layer", "unknown"),
                        "details": f.get("details", {}),
                    }
                    for f in findings
                ],
            )
            session.commit()
    except Exception as exc:  # noqa: BLE001
        logger.warning(
            "Failed to persist upload document %s: %s", doc.get("document_id"), exc
        )


def _execute_audit_job(
    job_id: str, file_ids: list[str], config_overrides: dict[str, Any]
) -> None:
    """Run the full audit pipeline for all specified files and update job state."""

    def _update(patch: dict[str, Any]) -> None:
        with _STORE_LOCK:
            _JOBS[job_id].update(patch)

    _update({"status": "running"})
    _record_mesh_job(job_id, len(file_ids), "executing")
    all_findings: list[dict[str, Any]] = []
    doc_manifests: list[dict[str, Any]] = []
    docs_processed = 0

    try:
        from oraculus_di_auditor.analysis import (
            analyze_document,
            find_blank_required_fields,
        )

        with _STORE_LOCK:
            files_snapshot = [_FILES[fid] for fid in file_ids if fid in _FILES]

        total = len(files_snapshot)
        # v2.9.3 D.2 — accumulate per-document blank-field state for the
        # corpus-scope rollup finding emitted after the per-document loop.
        blank_field_rollup: list[dict[str, Any]] = []

        for i, file_meta in enumerate(files_snapshot):
            _update(
                {
                    "progress": {
                        "phase": f"Analyzing {file_meta['name']} ({i + 1}/{total})",
                        "docs_processed": i,
                        "findings_count": len(all_findings),
                        "total_docs": total,
                    }
                }
            )

            path = Path(file_meta["path"])
            doc = ingest_uploaded_file(path)
            result = analyze_document(doc)
            findings = _flatten_findings(result, doc["document_id"])
            all_findings.extend(findings)
            docs_processed += 1

            # Persist document + analysis to DB so it appears in the
            # Documents, Anomalies, and Synthesis pages alongside
            # webhook-ingested documents.
            _persist_upload_document(
                doc, findings, config_overrides.get("jurisdiction")
            )

            manifest_entry: dict[str, Any] = {
                "document_id": doc["document_id"],
                "filename": file_meta["name"],
                "sha256": file_meta["sha256"],
                "size": file_meta["size"],
                "format": file_meta["format"],
                "finding_count": len(findings),
            }
            # v2.9.3 A.3 — surface OCR fallback status when available so
            # the evidence packet's executive summary can flag silent-
            # failure scans (PDFs where pypdf returned <500 chars and OCR
            # libs were absent).
            if doc.get("text_extraction"):
                manifest_entry["text_extraction"] = doc["text_extraction"]
            doc_manifests.append(manifest_entry)

            # v2.9.3 D.2 — accumulate blank-field state for the corpus
            # rollup; per-doc emission is gated off by default.
            blank = find_blank_required_fields(doc)
            if blank:
                blank_field_rollup.append(
                    {
                        "document_id": doc["document_id"],
                        "filename": file_meta["name"],
                        "missing": blank,
                    }
                )

        # v2.9.3 D.2 — single corpus-scope finding replacing the
        # per-document echoes that fired on 100% of corpora pre-2.9.3.
        # Severity stays MEDIUM because a corpus-wide pattern of
        # incomplete metadata is still a real records-management gap;
        # the difference is the count: 1 corpus finding instead of N.
        if blank_field_rollup:
            all_findings.append(
                {
                    "id": "admin:blank-required-fields-corpus",
                    "issue": (
                        f"Required metadata fields blank on "
                        f"{len(blank_field_rollup)} of "
                        f"{len(files_snapshot)} document(s)"
                    ),
                    "severity": "medium",
                    "layer": "administrative",
                    "scope": "corpus",
                    "document_id": "_corpus_",
                    "details": {
                        "total_affected": len(blank_field_rollup),
                        "total_documents": len(files_snapshot),
                        "affected_documents": blank_field_rollup,
                        "required_fields": [
                            "status",
                            "vote_result",
                            "meeting_date",
                            "agenda_number",
                        ],
                    },
                }
            )

        # Apply plain-language translations
        try:
            from oraculus_di_auditor.reporting.plain_language import translate_report

            all_findings = translate_report(all_findings)
        except Exception as exc:
            logger.warning("Plain-language translation failed: %s", exc)

        severity_counts: dict[str, int] = {
            "critical": 0,
            "high": 0,
            "medium": 0,
            "low": 0,
        }
        for f in all_findings:
            sev = f.get("severity", "low")
            if sev in severity_counts:
                severity_counts[sev] += 1

        _update(
            {
                "status": "complete",
                "results": {
                    "job_id": job_id,
                    "document_count": docs_processed,
                    "finding_count": len(all_findings),
                    "severity_summary": severity_counts,
                    "findings": all_findings,
                    "document_manifest": doc_manifests,
                    "generated_at": datetime.now(UTC).isoformat(),
                },
                "progress": {
                    "phase": "Complete",
                    "docs_processed": docs_processed,
                    "findings_count": len(all_findings),
                    "total_docs": total,
                },
            }
        )
        with _STORE_LOCK:
            completed_results = _JOBS[job_id].get("results")
        _record_mesh_job(
            job_id,
            len(file_ids),
            "completed",
            documents=docs_processed,
            findings=len(all_findings),
            severity_summary=severity_counts,
            results=completed_results,
        )

    except Exception as exc:
        logger.error("Audit job %s failed: %s", job_id, exc, exc_info=True)
        _update({"status": "error", "error": str(exc)})
        _record_mesh_job(job_id, len(file_ids), "failed", error=str(exc))


# ---------------------------------------------------------------------------
# Request model (defined at module level for FastAPI schema generation)
# ---------------------------------------------------------------------------


class _AuditRunRequest(BaseModel):  # type: ignore[misc]
    file_ids: list[str] = PydanticField(default_factory=list)
    jurisdiction: str | None = None
    agencies: list[str] | None = None


# ---------------------------------------------------------------------------
# Route registration
# ---------------------------------------------------------------------------


def register_upload_routes(app: Any) -> None:
    """Register all upload and audit endpoints on *app*.

    Safe to call when FastAPI is unavailable — silently does nothing.
    """
    if not _FASTAPI_AVAILABLE:
        return  # pragma: no cover

    router = APIRouter(tags=["upload", "audit"])

    # -- Upload endpoints -------------------------------------------------------

    @router.post("/api/v1/upload")
    async def upload_file(file: UploadFile = File(...)) -> dict[str, Any]:
        """Upload a single document (PDF, JSON, TXT, XML)."""
        ext = Path(file.filename or "").suffix.lower()
        if ext not in _ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Unsupported file type '{ext}'. "
                    f"Allowed: {', '.join(sorted(_ALLOWED_EXTENSIONS))}"
                ),
            )
        content = await file.read()
        file_id = str(uuid.uuid4())[:8]
        safe_name = Path(file.filename or "document").name
        dest = _UPLOAD_DIR / f"{file_id}_{safe_name}"
        dest.write_bytes(content)
        meta: dict[str, Any] = {
            "file_id": file_id,
            "name": safe_name,
            "size": len(content),
            "sha256": _sha256_bytes(content),
            "format": ext.lstrip("."),
            "path": str(dest),
            "uploaded_at": datetime.now(UTC).isoformat(),
        }
        with _STORE_LOCK:
            _FILES[file_id] = meta
        return meta

    @router.post("/api/v1/upload/batch")
    async def upload_batch(files: list[UploadFile] = File(...)) -> dict[str, Any]:
        """Upload multiple documents at once."""
        uploaded: list[dict[str, Any]] = []
        errors: list[dict[str, str]] = []
        for file in files:
            ext = Path(file.filename or "").suffix.lower()
            if ext not in _ALLOWED_EXTENSIONS:
                errors.append(
                    {"name": file.filename or "", "error": f"Unsupported type '{ext}'"}
                )
                continue
            content = await file.read()
            file_id = str(uuid.uuid4())[:8]
            safe_name = Path(file.filename or "document").name
            dest = _UPLOAD_DIR / f"{file_id}_{safe_name}"
            dest.write_bytes(content)
            meta = {
                "file_id": file_id,
                "name": safe_name,
                "size": len(content),
                "sha256": _sha256_bytes(content),
                "format": ext.lstrip("."),
                "path": str(dest),
                "uploaded_at": datetime.now(UTC).isoformat(),
            }
            with _STORE_LOCK:
                _FILES[file_id] = meta
            uploaded.append(meta)
        return {"uploaded": uploaded, "errors": errors}

    @router.post("/api/v1/upload/image")
    async def upload_image(file: UploadFile = File(...)) -> dict[str, Any]:
        """Upload a JPEG or PNG image and extract text via OCR.

        Returns the extracted text plus a file_id that can be used in
        POST /api/v1/audit/run just like any other uploaded file.
        """
        ext = Path(file.filename or "").suffix.lower()
        if ext not in _ALLOWED_IMAGE_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Unsupported image type '{ext}'. "
                    f"Allowed: {', '.join(sorted(_ALLOWED_IMAGE_EXTENSIONS))}"
                ),
            )
        content = await file.read()
        file_id = str(uuid.uuid4())[:8]
        safe_name = Path(file.filename or "image").name
        dest = _UPLOAD_DIR / f"{file_id}_{safe_name}"
        dest.write_bytes(content)

        raw_text, ocr_method = _ocr_image(dest)

        # Store as a TXT-format document so the audit pipeline can ingest it
        txt_path = _UPLOAD_DIR / f"{file_id}_{Path(safe_name).stem}.txt"
        txt_path.write_text(raw_text, encoding="utf-8")

        meta: dict[str, Any] = {
            "file_id": file_id,
            "name": safe_name,
            "size": len(content),
            "sha256": _sha256_bytes(content),
            "format": ext.lstrip("."),
            "path": str(txt_path),  # point pipeline at the extracted text
            "uploaded_at": datetime.now(UTC).isoformat(),
            "ocr_method": ocr_method,
        }
        with _STORE_LOCK:
            _FILES[file_id] = meta

        return {**meta, "raw_text": raw_text[:500] if raw_text else ""}

    @router.get("/api/v1/upload/files")
    async def list_files() -> dict[str, Any]:
        """List all files uploaded in this server session."""
        with _STORE_LOCK:
            files = list(_FILES.values())
        return {"files": files, "count": len(files)}

    @router.delete("/api/v1/upload/files/{file_id}")
    async def delete_file(file_id: str) -> dict[str, str]:
        """Remove an uploaded file by ID."""
        with _STORE_LOCK:
            if file_id not in _FILES:
                raise HTTPException(
                    status_code=404, detail=f"File '{file_id}' not found"
                )
            meta = _FILES.pop(file_id)
        path = Path(meta["path"])
        if path.exists():
            path.unlink()
        return {"status": "deleted", "file_id": file_id}

    # -- Audit endpoints --------------------------------------------------------

    @router.post("/api/v1/audit/run")
    async def run_audit(request: _AuditRunRequest) -> dict[str, Any]:
        """Start an audit job. Uses all uploaded files if file_ids is empty."""
        with _STORE_LOCK:
            available_ids = list(_FILES.keys())

        file_ids = request.file_ids if request.file_ids else available_ids
        if not file_ids:
            raise HTTPException(
                status_code=400,
                detail="No files available. Upload documents first.",
            )

        with _STORE_LOCK:
            missing = [fid for fid in file_ids if fid not in _FILES]
        if missing:
            raise HTTPException(status_code=404, detail=f"Files not found: {missing}")

        job_id = str(uuid.uuid4())
        with _STORE_LOCK:
            _JOBS[job_id] = {
                "job_id": job_id,
                "status": "pending",
                "file_ids": file_ids,
                "progress": {
                    "phase": "Queued",
                    "docs_processed": 0,
                    "findings_count": 0,
                    "total_docs": len(file_ids),
                },
                "results": None,
                "error": None,
                "created_at": datetime.now(UTC).isoformat(),
            }

        threading.Thread(
            target=_execute_audit_job,
            args=(job_id, file_ids, {"jurisdiction": request.jurisdiction}),
            daemon=True,
        ).start()

        return {"job_id": job_id, "status": "pending", "file_count": len(file_ids)}

    @router.get("/api/v1/audit/status/{job_id}")
    async def audit_status(job_id: str) -> dict[str, Any]:
        """Return current status and progress for a job."""
        with _STORE_LOCK:
            job = _JOBS.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail=f"Job '{job_id}' not found")
        return {
            "job_id": job_id,
            "status": job["status"],
            "progress": job["progress"],
            "error": job.get("error"),
        }

    @router.get("/api/v1/audit/results/{job_id}")
    async def audit_results(job_id: str) -> dict[str, Any]:
        """Return complete audit results. Falls back to DB when the in-memory
        job has been evicted (e.g. after a server restart)."""
        with _STORE_LOCK:
            job = _JOBS.get(job_id)
        if job:
            if job["status"] == "error":
                raise HTTPException(
                    status_code=500, detail=job.get("error", "Audit failed")
                )
            if job["status"] != "complete":
                return {"job_id": job_id, "status": job["status"], "results": None}
            return {"job_id": job_id, "status": "complete", "results": job["results"]}

        # Job not in memory — try DB fallback.
        try:
            from oraculus_di_auditor.db import models as db_models
            from oraculus_di_auditor.db.session import get_db

            with get_db() as session:
                row = (
                    session.query(db_models.MeshExecutionJob)
                    .filter(db_models.MeshExecutionJob.job_id == job_id)
                    .one_or_none()
                )
                if row and row.results_json:
                    return {
                        "job_id": job_id,
                        "status": "complete",
                        "results": json.loads(row.results_json),
                    }
        except Exception as exc:  # noqa: BLE001
            logger.warning("DB fallback for audit results failed: %s", exc)

        raise HTTPException(status_code=404, detail=f"Job '{job_id}' not found")

    @router.get("/api/v1/audit/history")
    async def audit_history(page: int = 1, per_page: int = 100) -> dict[str, Any]:
        """Paginated list of completed audit jobs from the DB, newest first.

        Returns lightweight summaries (no findings array) so the frontend
        history list can support thousands of entries without hitting
        localStorage size limits. Full results are fetched on demand via
        GET /api/v1/audit/results/{job_id}.
        """
        per_page = max(1, min(per_page, 500))
        try:
            from oraculus_di_auditor.db import models as db_models
            from oraculus_di_auditor.db.session import get_db

            with get_db() as session:
                base = (
                    session.query(db_models.MeshExecutionJob)
                    .filter(
                        db_models.MeshExecutionJob.job_type == "audit",
                        db_models.MeshExecutionJob.results_json.isnot(None),
                    )
                    .order_by(db_models.MeshExecutionJob.completed_at.desc())
                )
                total = base.count()
                rows = base.offset((page - 1) * per_page).limit(per_page).all()

                items = []
                for row in rows:
                    try:
                        r = json.loads(row.results_json or "{}")
                        manifest = r.get("document_manifest", [])
                        first_name = manifest[0]["filename"] if manifest else "Unknown"
                        doc_count = r.get("document_count", 0)
                        items.append(
                            {
                                "job_id": row.job_id,
                                "status": row.status,
                                "completed_at": (
                                    row.completed_at.isoformat()
                                    if row.completed_at
                                    else None
                                ),
                                "generated_at": r.get("generated_at"),
                                "document_count": doc_count,
                                "finding_count": r.get("finding_count", 0),
                                "severity_summary": r.get("severity_summary", {}),
                                "first_filename": first_name,
                                "more_docs": max(0, doc_count - 1),
                            }
                        )
                    except Exception:  # noqa: BLE001
                        continue

        except Exception as exc:  # noqa: BLE001
            logger.warning("audit_history query failed: %s", exc)
            return {
                "items": [],
                "total": 0,
                "page": page,
                "per_page": per_page,
                "has_more": False,
            }

        return {
            "items": items,
            "total": total,
            "page": page,
            "per_page": per_page,
            "has_more": (page * per_page) < total,
        }

    @router.get("/api/v1/audit/export/{job_id}")
    async def export_audit(job_id: str, format: str = "markdown") -> Any:  # noqa: A002
        """Download the audit report. format: markdown | html | pdf | docx"""
        with _STORE_LOCK:
            job = _JOBS.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail=f"Job '{job_id}' not found")
        if job["status"] != "complete":
            raise HTTPException(status_code=409, detail="Job not yet complete")

        report_md = _build_markdown_report(job["results"])
        short_id = job_id[:8]

        if format == "markdown":
            return StreamingResponse(
                io.BytesIO(report_md.encode("utf-8")),
                media_type="text/markdown",
                headers={
                    "Content-Disposition": f'attachment; filename="audit_report_{short_id}.md"'
                },
            )

        if format == "html":
            try:
                import markdown as md_lib  # type: ignore[import]

                body = md_lib.markdown(report_md, extensions=["tables"])
            except ImportError:
                body = f"<pre>{report_md}</pre>"
            html = f"<!DOCTYPE html><html><body>{body}</body></html>"
            return StreamingResponse(
                io.BytesIO(html.encode("utf-8")),
                media_type="text/html",
                headers={
                    "Content-Disposition": f'attachment; filename="audit_report_{short_id}.html"'
                },
            )

        if format in ("pdf", "docx"):
            try:
                from oraculus_di_auditor.reporting.format_converters import (
                    markdown_to_docx,
                    markdown_to_pdf,
                )

                if format == "pdf":
                    content = markdown_to_pdf(report_md)
                    mime = "application/pdf"
                else:
                    content = markdown_to_docx(report_md)
                    mime = (
                        "application/vnd.openxmlformats-officedocument"
                        ".wordprocessingml.document"
                    )
                return StreamingResponse(
                    io.BytesIO(content),
                    media_type=mime,
                    headers={
                        "Content-Disposition": (
                            f'attachment; filename="audit_report_{short_id}.{format}"'
                        )
                    },
                )
            except Exception as exc:
                raise HTTPException(
                    status_code=501,
                    detail=f"{format.upper()} export not available: {exc}",
                ) from exc

        raise HTTPException(
            status_code=400,
            detail=f"Unknown format '{format}'. Supported: markdown, html, pdf, docx",
        )

    @router.get("/api/v1/audit/evidence-packet/{job_id}")
    async def evidence_packet(job_id: str) -> Any:
        """Download a ZIP evidence packet containing the full audit materials."""
        with _STORE_LOCK:
            job = _JOBS.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail=f"Job '{job_id}' not found")
        if job["status"] != "complete":
            raise HTTPException(status_code=409, detail="Job not yet complete")

        try:
            from oraculus_di_auditor.reporting.evidence_packet import (
                generate_evidence_packet,
            )

            zip_bytes = generate_evidence_packet(job["results"])
        except Exception as exc:
            raise HTTPException(
                status_code=500, detail=f"Evidence packet error: {exc}"
            ) from exc

        return StreamingResponse(
            io.BytesIO(zip_bytes),
            media_type="application/zip",
            headers={
                "Content-Disposition": (
                    f'attachment; filename="evidence_packet_{job_id[:8]}.zip"'
                )
            },
        )

    app.include_router(router)

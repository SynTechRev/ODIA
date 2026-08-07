"""Tests for cards.tcams — T.C.A.M.S. corpus report generation."""

import pytest

sqlalchemy = pytest.importorskip("sqlalchemy")
docx = pytest.importorskip("docx")

from datetime import UTC, datetime  # noqa: E402

from sqlalchemy import create_engine  # noqa: E402
from sqlalchemy.orm import sessionmaker  # noqa: E402

from oraculus_di_auditor.cards.tcams import (  # noqa: E402
    _query_anchor_frequency,
    _query_axis_sums,
    _query_band_distribution,
    _query_detector_heatmap,
    _query_top_entities,
    build_tcams_report,
)
from oraculus_di_auditor.db.models import (  # noqa: E402
    Base,
    CasiScore,
    CommercialDocument,
    CommercialEntity,
    ContraFinding,
)

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def session(tmp_path):
    engine = create_engine(f"sqlite:///{tmp_path}/test_tcams.db")
    Base.metadata.create_all(engine)
    SessionFactory = sessionmaker(bind=engine)  # noqa: N806
    sess = SessionFactory()
    yield sess
    sess.close()
    engine.dispose()


def _make_entity(entity_id: str, name: str) -> CommercialEntity:
    return CommercialEntity(
        entity_id=entity_id,
        canonical_name=name,
        naics="5112",
        in_contra_corpus=True,
    )


def _make_doc(doc_hash: str, entity_id: str) -> CommercialDocument:
    return CommercialDocument(
        document_hash=doc_hash,
        entity_id=entity_id,
        doc_type="tos",
        retrieval_ts=datetime.now(UTC),
    )


def _make_score(doc_hash: str, **axes) -> CasiScore:
    rf = axes.get("remedy_foreclosure", 10)
    ded = axes.get("data_extraction_depth", 10)
    mc = axes.get("modification_and_consent", 10)
    pa = axes.get("procedural_adhesion", 10)
    eca = axes.get("enforcement_cost_asymmetry", 10)
    agg = rf + ded + mc + pa + eca
    return CasiScore(
        document_hash=doc_hash,
        remedy_foreclosure=rf,
        data_extraction_depth=ded,
        modification_and_consent=mc,
        procedural_adhesion=pa,
        enforcement_cost_asymmetry=eca,
        aggregate=agg,
        band=axes.get("band", "Elevated Asymmetry"),
        framework_version="1.0",
    )


def _make_finding(
    finding_id: str,
    doc_hash: str,
    layer: str = "L-11",
    sub_detector: str = "A",
    doctrinal_anchor: str = "arbitration clause",
    evidence_excerpt: str = "binding arbitration waiving class rights",
    severity: str = "high",
) -> ContraFinding:
    return ContraFinding(
        finding_id=finding_id,
        document_hash=doc_hash,
        layer=layer,
        sub_detector=sub_detector,
        severity=severity,
        doctrinal_anchor=doctrinal_anchor,
        evidence_excerpt=evidence_excerpt,
    )


@pytest.fixture()
def populated_session(session):
    e1 = _make_entity("eid-1", "TechCorp Inc.")
    e2 = _make_entity("eid-2", "DataCo LLC")
    session.add_all([e1, e2])

    d1 = _make_doc("abc" + "0" * 61, "eid-1")
    d2 = _make_doc("def" + "0" * 61, "eid-1")
    d3 = _make_doc("ghi" + "0" * 61, "eid-2")
    session.add_all([d1, d2, d3])

    s1 = _make_score(
        d1.document_hash,
        remedy_foreclosure=18,
        data_extraction_depth=16,
        modification_and_consent=14,
        procedural_adhesion=12,
        enforcement_cost_asymmetry=10,
        band="Severe Asymmetry",
    )
    s2 = _make_score(
        d2.document_hash,
        remedy_foreclosure=8,
        data_extraction_depth=8,
        modification_and_consent=8,
        procedural_adhesion=8,
        enforcement_cost_asymmetry=8,
        band="Elevated Asymmetry",
    )
    s3 = _make_score(
        d3.document_hash,
        remedy_foreclosure=5,
        data_extraction_depth=4,
        modification_and_consent=3,
        procedural_adhesion=3,
        enforcement_cost_asymmetry=3,
        band="Baseline Adhesion",
    )
    session.add_all([s1, s2, s3])

    f1 = _make_finding(
        "f-001",
        d1.document_hash,
        layer="L-11",
        sub_detector="A",
        doctrinal_anchor="forced arbitration",
    )
    f2 = _make_finding(
        "f-002",
        d1.document_hash,
        layer="L-12",
        sub_detector="B",
        doctrinal_anchor="class action waiver",
    )
    f3 = _make_finding(
        "f-003",
        d2.document_hash,
        layer="L-11",
        sub_detector="A",
        doctrinal_anchor="forced arbitration",
    )
    f4 = _make_finding(
        "f-004",
        d3.document_hash,
        layer="L-13",
        sub_detector="C",
        doctrinal_anchor="data monetization",
    )
    session.add_all([f1, f2, f3, f4])
    session.commit()
    return session


# ---------------------------------------------------------------------------
# Query unit tests
# ---------------------------------------------------------------------------


class TestQueryBandDistribution:
    def test_empty_db(self, session):
        dist = _query_band_distribution(session)
        assert dist == {
            "Baseline Adhesion": 0,
            "Elevated Asymmetry": 0,
            "Substantial Asymmetry": 0,
            "Severe Asymmetry": 0,
            "Foreclosure Regime": 0,
        }

    def test_counts_correct(self, populated_session):
        dist = _query_band_distribution(populated_session)
        assert dist["Severe Asymmetry"] == 1
        assert dist["Elevated Asymmetry"] == 1
        assert dist["Baseline Adhesion"] == 1

    def test_all_bands_present_in_output(self, populated_session):
        dist = _query_band_distribution(populated_session)
        from oraculus_di_auditor.cards.tcams import _BAND_ORDER

        for band in _BAND_ORDER:
            assert band in dist


class TestQueryTopEntities:
    def test_empty_db(self, session):
        result = _query_top_entities(session)
        assert result == []

    def test_ordered_by_max_score(self, populated_session):
        result = _query_top_entities(populated_session)
        scores = [row[1] for row in result]
        assert scores == sorted(scores, reverse=True)

    def test_eid1_has_two_docs(self, populated_session):
        result = _query_top_entities(populated_session)
        eid1_row = next(r for r in result if r[0] == "TechCorp Inc.")
        assert eid1_row[3] == 2  # doc_count


class TestQueryAxisSums:
    def test_empty_db(self, session):
        data = _query_axis_sums(session)
        assert data["doc_count"] == 0
        assert data["remedy_foreclosure"] == 0

    def test_sums_correct(self, populated_session):
        data = _query_axis_sums(populated_session)
        assert data["doc_count"] == 3
        # s1.remedy_foreclosure=18, s2=8, s3=5
        assert data["remedy_foreclosure"] == 31


class TestQueryAnchorFrequency:
    def test_empty_db(self, session):
        result = _query_anchor_frequency(session)
        assert result == []

    def test_forced_arbitration_appears_twice(self, populated_session):
        result = _query_anchor_frequency(populated_session)
        # forced arbitration appears in d1 and d2 (2 unique documents)
        fa = next((r for r in result if r[0] == "forced arbitration"), None)
        assert fa is not None
        assert fa[1] == 2  # doc_count
        assert fa[2] == 2  # finding_count

    def test_ordered_by_doc_count_desc(self, populated_session):
        result = _query_anchor_frequency(populated_session)
        doc_counts = [r[1] for r in result]
        assert doc_counts == sorted(doc_counts, reverse=True)


class TestQueryDetectorHeatmap:
    def test_empty_db(self, session):
        result = _query_detector_heatmap(session)
        assert result == []

    def test_l11a_is_top_detector(self, populated_session):
        result = _query_detector_heatmap(populated_session)
        top = result[0]
        assert top[0] == "L-11"
        assert top[1] == "A"
        assert top[2] == 2

    def test_ordered_by_count_desc(self, populated_session):
        result = _query_detector_heatmap(populated_session)
        counts = [r[2] for r in result]
        assert counts == sorted(counts, reverse=True)


# ---------------------------------------------------------------------------
# Integration: build_tcams_report
# ---------------------------------------------------------------------------


class TestBuildTcamsReport:
    def test_returns_path(self, populated_session, tmp_path):
        path = build_tcams_report(populated_session, tmp_path)
        assert path.exists()
        assert path.suffix == ".docx"

    def test_creates_output_dir(self, populated_session, tmp_path):
        out = tmp_path / "reports" / "tcams"
        path = build_tcams_report(populated_session, out)
        assert out.exists()
        assert path.parent == out

    def test_empty_db_produces_valid_docx(self, session, tmp_path):
        path = build_tcams_report(session, tmp_path)
        assert path.exists()
        assert path.stat().st_size > 0

    def test_filename_includes_tcams(self, populated_session, tmp_path):
        path = build_tcams_report(populated_session, tmp_path)
        assert "tcams_report" in path.name

    def test_docx_contains_section_headings(self, populated_session, tmp_path):
        from docx import Document as DocxDoc

        path = build_tcams_report(populated_session, tmp_path)
        doc = DocxDoc(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "CASI SCORE DISTRIBUTION" in text
        assert "TOP-10 ENTITIES" in text
        assert "AXIS BREAKDOWN" in text
        assert "DOCTRINAL ANCHOR FREQUENCY" in text
        assert "L-DETECTOR HEATMAP" in text

    def test_generated_by_in_docx(self, populated_session, tmp_path):
        from docx import Document as DocxDoc

        path = build_tcams_report(populated_session, tmp_path, generated_by="TestSuite")
        doc = DocxDoc(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "TestSuite" in text

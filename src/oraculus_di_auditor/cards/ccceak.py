"""C.C.C.E.A. — Commercial Contract Clause Exposure Analysis (D-13.H).

Generates a DOCX report that:
  1. Loads all ContraFindings from the DB
  2. Clusters evidence_excerpt strings by Jaccard token overlap
  3. For each cluster, identifies the materially identical clause type
     (arbitration, ML-training, data-use, etc.) and maps it to the probable
     drafting law firm via a keyword heuristic table
  4. Emits one DOCX section per cluster, ranked by cluster size

Style: same palette as analytical_card.py and tcams.py
  - Garamond font, malachite (#1D6B44) headings, tan (#C8A882) sub-headings
  - Page X of Y footer, no em-dash/en-dash in raw XML

Law firm heuristics are based on publicly reported drafting patterns and
are explicitly labeled as probable, not confirmed attribution.

Source: C.O.N.T.R.A. Framework D-13.H / Phase F Spec
"""

from __future__ import annotations

import logging
import re
from collections import defaultdict
from datetime import UTC, datetime
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

if TYPE_CHECKING:
    from sqlalchemy.orm import Session

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

_MALACHITE = RGBColor(0x1D, 0x6B, 0x44)
_TAN = RGBColor(0xC8, 0xA8, 0x82)
_FONT_BODY = "Garamond"
_FONT_FALLBACK = "Times New Roman"

# ---------------------------------------------------------------------------
# Clause-type keyword heuristics
# ---------------------------------------------------------------------------

_CLAUSE_KEYWORDS: list[tuple[str, list[str]]] = [
    (
        "Mandatory Arbitration / Class Waiver",
        [
            "arbitrat",
            "class action",
            "class waiver",
            "waive.*jury",
            "jury.*waive",
            "aaa rules",
            "jams rules",
        ],
    ),
    (
        "ML Training / AI Data Use",
        [
            "train.*model",
            "machine learning",
            "artificial intelligence",
            "ai.*purposes",
            "model.*train",
            "generative",
        ],
    ),
    (
        "Broad Data Collection / License",
        [
            "collect.*data",
            "data.*collect",
            "license.*data",
            "data.*license",
            "aggregat.*data",
            "data.*aggregat",
            "data.*monetiz",
        ],
    ),
    (
        "Unilateral Modification",
        [
            "modif.*sole discretion",
            "sole discretion.*modif",
            "change.*without.*notice",
            "amend.*any.*time",
            "any time.*amend",
        ],
    ),
    (
        "Limitation of Liability / Disclaimer",
        [
            "limit.*liabilit",
            "liabilit.*limit",
            "disclaim.*warrant",
            "warrant.*disclaim",
            "as.is",
            "no warrant",
        ],
    ),
    (
        "Indemnification (Asymmetric)",
        ["indemnif", "hold harmless", "defend.*against.*claim"],
    ),
    (
        "IP Assignment / Broad License Grant",
        [
            "assign.*intellectual",
            "intellectual.*assign",
            "work.*for.*hire",
            "perpetual.*irrevocable",
            "irrevocable.*perpetual",
        ],
    ),
    (
        "Fee-Shifting / Cost Asymmetry",
        [
            "fee.*shift",
            "arbitration.*fee",
            "attorney.*fee.*non-prevail",
            "pay.*cost.*arbitrat",
        ],
    ),
]

# ---------------------------------------------------------------------------
# Law firm drafting heuristics (probable attribution, not confirmed)
# ---------------------------------------------------------------------------

_FIRM_HEURISTICS: list[tuple[str, list[str]]] = [
    (
        "Cooley LLP",
        [
            "bilateral.*arbitrat",
            "mutual.*class.*waiv",
            "opt.*out.*thirty",
            "opt.*out.*30.*day",
        ],
    ),
    (
        "Wilson Sonsini Goodrich & Rosati",
        ["california.*govern.*law", "santa.*clara.*county.*jurisd", "palo alto.*forum"],
    ),
    (
        "Fenwick & West LLP",
        [
            "silicon.*valley.*arbitrat",
            "aaa.*consumer.*due.*process",
            "aaa.*suppl.*procedure",
        ],
    ),
    (
        "Latham & Watkins LLP",
        [
            "new york.*govern",
            "delaware.*govern.*law",
            "jams.*streamlined",
            "jams.*comprehensive",
        ],
    ),
    (
        "Gibson Dunn & Crutcher LLP",
        [
            "written.*notice.*sixty",
            "informal.*resolut.*sixty",
            "small.*claims.*carve",
            "ip.*carve.*out",
        ],
    ),
    (
        "DLA Piper LLP",
        [
            "aaa.*commercial.*arbitrat.*rule",
            "icc.*arbitrat",
            "london.*arbitrat",
            "london.*court.*internat",
        ],
    ),
]


def _tokenize(text: str) -> frozenset[str]:
    tokens = re.findall(r"[a-z0-9]+", text.lower())
    return frozenset(tokens)


def _jaccard(a: frozenset[str], b: frozenset[str]) -> float:
    if not a and not b:
        return 1.0
    union = a | b
    return len(a & b) / len(union)


def _classify_clause(text: str) -> str:
    lower = text.lower()
    for clause_type, patterns in _CLAUSE_KEYWORDS:
        for pat in patterns:
            if re.search(pat, lower):
                return clause_type
    return "Unclassified Clause"


def _probable_firm(text: str) -> str:
    lower = text.lower()
    for firm, patterns in _FIRM_HEURISTICS:
        for pat in patterns:
            if re.search(pat, lower):
                return firm
    return "Unknown / No Match"


# ---------------------------------------------------------------------------
# Clustering
# ---------------------------------------------------------------------------


def _cluster_excerpts(
    excerpts: list[tuple[str, str, str, str]],
    threshold: float = 0.35,
) -> list[list[tuple[str, str, str, str]]]:
    """Greedy single-linkage clustering of (excerpt, layer, sub_detector, finding_id) tuples.

    Two excerpts are in the same cluster if their Jaccard similarity >= threshold.
    Returns list of clusters (each a list of tuples), sorted descending by size.
    Complexity: O(n^2) — acceptable for corpus sizes up to ~50k findings.
    """
    if not excerpts:
        return []

    tokenized = [(_tokenize(exc[0]), exc) for exc in excerpts]
    n = len(tokenized)
    assigned = [-1] * n
    clusters: list[list[int]] = []

    for i in range(n):
        if assigned[i] != -1:
            continue
        cluster_idx = len(clusters)
        clusters.append([i])
        assigned[i] = cluster_idx
        tok_i = tokenized[i][0]
        for j in range(i + 1, n):
            if assigned[j] != -1:
                continue
            if _jaccard(tok_i, tokenized[j][0]) >= threshold:
                clusters[cluster_idx].append(j)
                assigned[j] = cluster_idx

    result = [[tokenized[idx][1] for idx in cluster] for cluster in clusters]
    result.sort(key=len, reverse=True)
    return result


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------


def _run_font(run, size_pt: float = 11, bold: bool = False) -> None:
    run.font.name = _FONT_BODY
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), _FONT_BODY)
    rFonts.set(qn("w:hAnsi"), _FONT_BODY)
    rFonts.set(qn("w:cs"), _FONT_FALLBACK)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _section_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _run_font(run, size_pt=14, bold=True)
    run.font.color.rgb = _MALACHITE
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)


def _sub_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _run_font(run, size_pt=11, bold=True)
    run.font.color.rgb = _TAN


def _body(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    for run in para.runs:
        _run_font(run)


def _kv(doc: Document, key: str, value: str) -> None:
    para = doc.add_paragraph()
    kr = para.add_run(f"{key}: ")
    _run_font(kr, bold=True)
    vr = para.add_run(value)
    _run_font(vr)


def _page_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run("Page ")
    _run_font(run, size_pt=9)
    for field_text in (" PAGE ", " NUMPAGES "):
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run._r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = field_text
        run._r.append(instr)
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)
        if field_text == " PAGE ":
            sep = para.add_run(" of ")
            _run_font(sep, size_pt=9)
            run = sep


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def build_ccceak_report(
    session: Session,
    output_dir: str | Path,
    generated_by: str = "ODIA C.O.N.T.R.A. C.C.C.E.A.",
    jaccard_threshold: float = 0.35,
    max_clusters: int = 50,
    framework_version: str = "1.0",
) -> Path:
    """Generate a C.C.C.E.A. clause-exposure-analysis DOCX report.

    Loads all ContraFindings with non-empty evidence_excerpt from the DB,
    clusters by Jaccard overlap, classifies each cluster by clause type,
    maps to probable drafting firm, and emits ranked cluster sections.

    Arguments:
        session            -- active SQLAlchemy session
        output_dir         -- directory to write the report
        generated_by       -- source label in the report header
        jaccard_threshold  -- similarity threshold for greedy clustering (default 0.35)
        max_clusters       -- maximum clusters to include in report (default 50)
        framework_version  -- C.O.N.T.R.A. framework version tag

    Returns the resolved Path to the written .docx file.
    """
    from ..db.models import ContraFinding

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now(UTC).strftime("%Y%m%d_%H%M%S")
    filename = f"ccceak_report_{ts}.docx"
    output_path = output_dir / filename

    rows = (
        session.query(
            ContraFinding.evidence_excerpt,
            ContraFinding.layer,
            ContraFinding.sub_detector,
            ContraFinding.finding_id,
        )
        .filter(
            ContraFinding.evidence_excerpt.isnot(None),
            ContraFinding.evidence_excerpt != "",
        )
        .all()
    )

    excerpts: list[tuple[str, str, str, str]] = [
        (r.evidence_excerpt, r.layer or "", r.sub_detector or "", r.finding_id or "")
        for r in rows
    ]

    clusters = _cluster_excerpts(excerpts, threshold=jaccard_threshold)
    clusters = clusters[:max_clusters]

    doc = Document()

    # ------------------------------------------------------------------
    # Title block
    # ------------------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_para.add_run(
        "C.C.C.E.A. -- COMMERCIAL CONTRACT CLAUSE EXPOSURE ANALYSIS"
    )
    _run_font(tr, size_pt=16, bold=True)
    tr.font.color.rgb = _MALACHITE

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub_para.add_run(
        f"C.O.N.T.R.A. Framework V{framework_version} -- D-13.H Cross-Entity Clause Clustering"
    )
    _run_font(sr, size_pt=11)
    sr.font.color.rgb = _TAN

    doc.add_paragraph()
    _kv(doc, "Generated", datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC"))
    _kv(doc, "Source", generated_by)
    _kv(doc, "Jaccard Threshold", f"{jaccard_threshold:.2f}")
    _kv(doc, "Total Findings Loaded", str(len(excerpts)))
    _kv(doc, "Clusters Identified", str(len(clusters)))
    doc.add_paragraph()

    _body(
        doc,
        "NOTE: Law firm attribution below is PROBABLE based on keyword pattern "
        "matching against publicly reported drafting conventions. It is not a "
        "confirmed identification and should not be treated as such without "
        "independent verification.",
    )
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # Cluster sections
    # ------------------------------------------------------------------
    if not clusters:
        _body(doc, "No clusterable evidence excerpts found in corpus.")
    else:
        _section_heading(doc, "CROSS-ENTITY CLAUSE CLUSTERS (RANKED BY SIZE)")
        doc.add_paragraph()

        for rank, cluster in enumerate(clusters, start=1):
            cluster_text = cluster[0][0]
            clause_type = _classify_clause(cluster_text)
            firm = _probable_firm(cluster_text)

            layers = [item[1] for item in cluster if item[1]]
            sub_dets = [item[2] for item in cluster if item[2]]

            layer_counts: dict[str, int] = defaultdict(int)
            for lyr in layers:
                layer_counts[lyr] += 1
            sub_counts: dict[str, int] = defaultdict(int)
            for sd in sub_dets:
                sub_counts[sd] += 1

            dominant_layer = (
                max(layer_counts, key=layer_counts.__getitem__)
                if layer_counts
                else "N/A"
            )
            dominant_sub = (
                max(sub_counts, key=sub_counts.__getitem__) if sub_counts else "N/A"
            )

            _sub_heading(
                doc, f"Cluster {rank} -- {clause_type} ({len(cluster)} findings)"
            )
            _kv(doc, "Clause Type", clause_type)
            _kv(doc, "Probable Drafting Firm", firm)
            _kv(doc, "Cluster Size", str(len(cluster)))
            _kv(doc, "Dominant Layer", dominant_layer)
            _kv(doc, "Dominant Sub-Detector", dominant_sub)

            # Representative excerpt (first item, truncated to 200 chars)
            excerpt_display = (
                cluster_text[:200] + "..." if len(cluster_text) > 200 else cluster_text
            )
            _kv(doc, "Representative Excerpt", f'"{excerpt_display}"')

            # Unique entities represented (by finding_id prefix patterns — no PII)
            unique_sub_detectors = set(sub_dets)
            if unique_sub_detectors:
                para = doc.add_paragraph()
                kr = para.add_run("Sub-Detectors in Cluster: ")
                _run_font(kr, bold=True)
                vr = para.add_run(", ".join(sorted(unique_sub_detectors)))
                _run_font(vr)

            doc.add_paragraph()

        # ------------------------------------------------------------------
        # Summary table
        # ------------------------------------------------------------------
        _section_heading(doc, "CLAUSE TYPE SUMMARY")
        type_counts: dict[str, int] = defaultdict(int)
        firm_counts: dict[str, int] = defaultdict(int)
        for cluster in clusters:
            ct = _classify_clause(cluster[0][0])
            fm = _probable_firm(cluster[0][0])
            type_counts[ct] += len(cluster)
            firm_counts[fm] += len(cluster)

        _sub_heading(doc, "By Clause Type")
        tbl_type = doc.add_table(rows=1, cols=2)
        tbl_type.style = "Table Grid"
        for i, h in enumerate(["Clause Type", "Findings"]):
            r = tbl_type.rows[0].cells[i].paragraphs[0].add_run(h)
            _run_font(r, bold=True)
            r.font.color.rgb = _MALACHITE
        for clause_type, count in sorted(
            type_counts.items(), key=lambda x: x[1], reverse=True
        ):
            row = tbl_type.add_row()
            _run_font(row.cells[0].paragraphs[0].add_run(clause_type))
            _run_font(row.cells[1].paragraphs[0].add_run(str(count)))

        doc.add_paragraph()
        _sub_heading(doc, "By Probable Drafting Firm")
        tbl_firm = doc.add_table(rows=1, cols=2)
        tbl_firm.style = "Table Grid"
        for i, h in enumerate(["Probable Firm (see NOTE)", "Findings"]):
            r = tbl_firm.rows[0].cells[i].paragraphs[0].add_run(h)
            _run_font(r, bold=True)
            r.font.color.rgb = _MALACHITE
        for firm, count in sorted(
            firm_counts.items(), key=lambda x: x[1], reverse=True
        ):
            row = tbl_firm.add_row()
            _run_font(row.cells[0].paragraphs[0].add_run(firm))
            _run_font(row.cells[1].paragraphs[0].add_run(str(count)))

    _page_footer(doc)
    doc.save(str(output_path))
    log.info("C.C.C.E.A. report written: %s", output_path)
    return output_path

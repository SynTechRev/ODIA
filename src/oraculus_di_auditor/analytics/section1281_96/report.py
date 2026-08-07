"""CCP § 1281.96 Summary Statistics Report generator.

Produces a .docx report containing:
    - Dataset overview (total cases, providers, date range)
    - Consumer prevailing rate table stratified by representation + claim tier
    - Arbitrator repeat-player concentration summary
    - Corporate repeat-player concentration table (top 10)
    - CONTRA corpus entity cross-reference table
    - Quality flag summary

python-docx is an optional dependency; importing this module raises ImportError
if it is not installed.  Install with: pip install python-docx

Usage:
    from oraculus_di_auditor.analytics.section1281_96.report import build_summary_report
    from oraculus_di_auditor.analytics.section1281_96 import (
        prevailing_rate_stratified,
        arbitrator_repeat_player_concentration,
        corporate_repeat_player_concentration,
    )

    cases = [...]   # list[NormalizedCase]
    build_summary_report(cases, path="report_q1_2024.docx")
"""

from __future__ import annotations

import logging
from collections import Counter
from datetime import datetime
from pathlib import Path

from .compute import (
    arbitrator_repeat_player_concentration,
    contra_corpus_entity_slice,
    corporate_repeat_player_concentration,
    prevailing_rate_stratified,
)
from .normalize import NormalizedCase

log = logging.getLogger(__name__)

_TIER_ORDER = ["UNDER_1K", "1K_10K", "10K_75K", "75K_300K", "OVER_300K", "UNKNOWN"]
_REP_ORDER = ["YES", "NO", "UNKNOWN"]


def _try_docx():
    try:
        import docx

        return docx
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for report generation: pip install python-docx"
        ) from exc


def _heading(doc, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _para(doc, text: str):
    doc.add_paragraph(text)


def _kv(doc, key: str, value: str | int | float):
    p = doc.add_paragraph()
    run = p.add_run(f"{key}: ")
    run.bold = True
    p.add_run(str(value))


def build_summary_report(
    cases: list[NormalizedCase],
    path: str | Path = "section_1281_96_report.docx",
    entity_ids: set[str] | None = None,
    generated_by: str = "ODIA § 1281.96 Pipeline",
) -> Path:
    """Build a .docx summary statistics report for the provided NormalizedCase list.

    Arguments:
        cases        -- list of NormalizedCase records (all providers combined)
        path         -- output file path (default: section_1281_96_report.docx)
        entity_ids   -- optional set of entity_ids where in_contra_corpus=True;
                        when provided, a CONTRA cross-reference section is included
        generated_by -- source label in the report header

    Returns the resolved Path of the written file.
    """
    docx = _try_docx()
    doc = docx.Document()
    out = Path(path)

    # ------------------------------------------------------------------ header
    doc.add_heading("CCP § 1281.96 — Consumer Arbitration Statistics Report", 0)
    _para(doc, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M UTC')}")
    _para(doc, f"Source: {generated_by}")
    doc.add_paragraph()

    # ---------------------------------------------------------------- overview
    _heading(doc, "1. Dataset Overview")
    provider_counts = Counter(c.provider for c in cases)
    years = sorted({c.case_year for c in cases})
    quality_counts: Counter[str] = Counter()
    for c in cases:
        for f in c.quality_flags:
            quality_counts[f] += 1

    _kv(doc, "Total Cases", len(cases))
    _kv(doc, "Providers", ", ".join(sorted(provider_counts)))
    _kv(doc, "Year Range", f"{min(years)} – {max(years)}" if years else "N/A")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Provider"
    tbl.rows[0].cells[1].text = "Cases"
    for prov, cnt in sorted(provider_counts.items()):
        row = tbl.add_row()
        row.cells[0].text = prov
        row.cells[1].text = str(cnt)

    doc.add_paragraph()

    # ------------------------------------------------ prevailing rate stratified
    _heading(doc, "2. Consumer Prevailing Rate (AWARD_AFTER_HEARING Only)")
    _para(
        doc,
        "Table denominates exclusively on cases where disposition_type = AWARD_AFTER_HEARING. "
        "95% Wilson confidence intervals shown in parentheses.",
    )

    strat = prevailing_rate_stratified(cases)
    tbl2 = doc.add_table(rows=1, cols=6)
    tbl2.style = "Table Grid"
    hdr = tbl2.rows[0].cells
    for i, label in enumerate(
        [
            "Represented",
            "Claim Tier",
            "N (hearings)",
            "Consumer Wins",
            "Win Rate",
            "95% CI",
        ]
    ):
        hdr[i].text = label

    for rep in _REP_ORDER:
        if rep not in strat:
            continue
        tier_data = strat[rep]
        for tier in _TIER_ORDER:
            if tier not in tier_data:
                continue
            d = tier_data[tier]
            row = tbl2.add_row()
            row.cells[0].text = rep
            row.cells[1].text = tier
            row.cells[2].text = str(d["n_cases"])
            row.cells[3].text = str(d["n_consumer_wins"])
            row.cells[4].text = f"{d['rate']:.1%}"
            row.cells[5].text = f"({d['ci_lower']:.1%} – {d['ci_upper']:.1%})"

    doc.add_paragraph()

    # ---------------------------------------- arbitrator repeat-player
    _heading(doc, "3. Arbitrator Repeat-Player Concentration")
    arb = arbitrator_repeat_player_concentration(cases)

    _kv(doc, "Total Case-Arbitrator Assignments", arb["total_case_assignments"])
    _kv(doc, "Unique Arbitrators", arb["unique_arbitrators"])
    _kv(
        doc,
        "Top 5% arbitrators account for",
        f"{arb['top_5pct_case_share']:.1%} of volume "
        f"({arb['top_5pct_arbitrators']} arbitrators)",
    )
    _kv(
        doc,
        "Top 10% arbitrators account for",
        f"{arb['top_10pct_case_share']:.1%} of volume "
        f"({arb['top_10pct_arbitrators']} arbitrators)",
    )
    _kv(
        doc,
        "Top 25% arbitrators account for",
        f"{arb['top_25pct_case_share']:.1%} of volume "
        f"({arb['top_25pct_arbitrators']} arbitrators)",
    )

    if arb["top_10_by_volume"]:
        _para(doc, "Top 10 arbitrators by case volume:")
        tbl3 = doc.add_table(rows=1, cols=3)
        tbl3.style = "Table Grid"
        h3 = tbl3.rows[0].cells
        h3[0].text = "Arbitrator"
        h3[1].text = "Cases"
        h3[2].text = "Share"
        for entry in arb["top_10_by_volume"]:
            row = tbl3.add_row()
            row.cells[0].text = entry["name"]
            row.cells[1].text = str(entry["case_count"])
            row.cells[2].text = f"{entry['share']:.1%}"

    doc.add_paragraph()

    # ---------------------------------------- corporate repeat-player
    _heading(doc, "4. Corporate Repeat-Player Concentration")
    corp = corporate_repeat_player_concentration(cases)

    _kv(doc, "Total Cases", corp["total_cases"])
    _kv(doc, "Unique Companies", corp["unique_companies"])
    _kv(doc, "Herfindahl-Hirschman Index (HHI)", f"{corp['herfindahl_index']:.0f}")
    _para(
        doc,
        "HHI interpretation: <1,500 = competitive; 1,500–2,500 = moderately concentrated; "
        ">2,500 = highly concentrated.",
    )

    if corp["top_10_by_volume"]:
        tbl4 = doc.add_table(rows=1, cols=5)
        tbl4.style = "Table Grid"
        h4 = tbl4.rows[0].cells
        for i, label in enumerate(
            ["Company", "Cases", "Share", "Award Cases", "Consumer Win Rate"]
        ):
            h4[i].text = label
        for entry in corp["top_10_by_volume"]:
            row = tbl4.add_row()
            row.cells[0].text = entry["name"]
            row.cells[1].text = str(entry["case_count"])
            row.cells[2].text = f"{entry['share']:.1%}"
            row.cells[3].text = str(entry["award_cases"])
            row.cells[4].text = (
                f"{entry['consumer_win_rate']:.1%}"
                if entry["consumer_win_rate"] is not None
                else "N/A"
            )

    doc.add_paragraph()

    # ---------------------------------------- CONTRA cross-reference
    if entity_ids:
        _heading(doc, "5. CONTRA Corpus Entity Cross-Reference")
        _para(
            doc,
            "Cases where the non-consumer party matches a CONTRA corpus entity "
            "(CommercialEntity.in_contra_corpus = True).",
        )
        sliced = contra_corpus_entity_slice(cases, entity_ids)
        if sliced:
            tbl5 = doc.add_table(rows=1, cols=4)
            tbl5.style = "Table Grid"
            h5 = tbl5.rows[0].cells
            for i, label in enumerate(
                ["Entity ID", "Cases", "Award Cases", "Consumer Win Rate"]
            ):
                h5[i].text = label
            for eid, ent_cases in sorted(sliced.items(), key=lambda kv: -len(kv[1])):
                award = [
                    c for c in ent_cases if c.disposition_type == "AWARD_AFTER_HEARING"
                ]
                wins = sum(1 for c in award if c.prevailing_party == "CONSUMER")
                rate = f"{wins / len(award):.1%}" if award else "N/A"
                row = tbl5.add_row()
                row.cells[0].text = eid
                row.cells[1].text = str(len(ent_cases))
                row.cells[2].text = str(len(award))
                row.cells[3].text = rate
        else:
            _para(doc, "No CONTRA corpus entities found in this dataset.")
        doc.add_paragraph()

    # ---------------------------------------- quality flags
    section_num = 6 if entity_ids else 5
    _heading(doc, f"{section_num}. Data Quality Flags")
    if quality_counts:
        tbl6 = doc.add_table(rows=1, cols=3)
        tbl6.style = "Table Grid"
        h6 = tbl6.rows[0].cells
        h6[0].text = "Flag"
        h6[1].text = "Count"
        h6[2].text = "Rate"
        for flag, cnt in sorted(quality_counts.items(), key=lambda kv: -kv[1]):
            row = tbl6.add_row()
            row.cells[0].text = flag
            row.cells[1].text = str(cnt)
            row.cells[2].text = f"{cnt / len(cases):.1%}" if cases else "N/A"
    else:
        _para(doc, "No quality flags raised.")

    doc.save(str(out))
    log.info("Report written: %s", out)
    return out

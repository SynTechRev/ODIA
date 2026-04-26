"""Format conversion utilities for ODIA audit reports.

Converts rendered Markdown reports to PDF, DOCX, and HTML for professional
distribution.  All converters degrade gracefully: if the required external
tool is unavailable the function returns None and logs a warning rather than
raising an exception.

Conversion priority:
  HTML  — pandoc → Python ``markdown`` library
  PDF   — pandoc → weasyprint → wkhtmltopdf
  DOCX  — pandoc only
"""

from __future__ import annotations

import json
import logging
import re
import shutil
import subprocess
from pathlib import Path
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:
    from oraculus_di_auditor.reporting.models import AuditReport

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Low-level converters
# ---------------------------------------------------------------------------


def markdown_to_html(markdown_text: str) -> str:
    """Convert Markdown to HTML.

    Uses pandoc if available, falls back to the Python ``markdown`` library.

    Args:
        markdown_text: Markdown source string.

    Returns:
        HTML string.  Never raises — falls back to a ``<pre>``-wrapped
        representation if neither pandoc nor the markdown library is present.
    """
    # Try pandoc first.
    if shutil.which("pandoc"):
        try:
            result = subprocess.run(
                ["pandoc", "--from=markdown", "--to=html"],
                input=markdown_text,
                capture_output=True,
                text=True,
                timeout=30,
                check=True,
            )
            return result.stdout
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            logger.warning("pandoc failed during HTML conversion: %s", exc)

    # Try Python markdown library.
    try:
        import markdown as md_lib  # type: ignore[import-untyped]

        return md_lib.markdown(markdown_text, extensions=["tables", "fenced_code"])
    except ImportError:
        logger.warning(
            "Neither pandoc nor the Python 'markdown' library is available. "
            "Returning pre-wrapped HTML fallback."
        )

    escaped = markdown_text.replace("&", "&amp;").replace("<", "&lt;")
    return f"<pre>{escaped}</pre>"


def markdown_to_pdf(
    markdown_text: str,
    output_path: Path | str,
    title: str = "ODIA Audit Report",
) -> Path | None:
    """Convert Markdown to PDF.

    Tries pandoc first, then weasyprint, then wkhtmltopdf.

    Args:
        markdown_text: Markdown source string.
        output_path: Destination ``.pdf`` file path.
        title: Document title embedded in the PDF metadata.

    Returns:
        Resolved Path of the written file on success, ``None`` if no PDF
        converter is available.
    """
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    # 1. pandoc
    if shutil.which("pandoc"):
        try:
            subprocess.run(
                [
                    "pandoc",
                    "--from=markdown",
                    "--to=pdf",
                    f"--metadata=title:{title}",
                    f"--output={out}",
                ],
                input=markdown_text,
                capture_output=True,
                text=True,
                timeout=60,
                check=True,
            )
            logger.info("PDF written via pandoc: %s", out)
            return out.resolve()
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            logger.warning("pandoc PDF conversion failed: %s", exc)

    # 2. weasyprint
    try:
        import weasyprint  # type: ignore[import-untyped]

        html = markdown_to_html(markdown_text)
        weasyprint.HTML(string=html).write_pdf(str(out))
        logger.info("PDF written via weasyprint: %s", out)
        return out.resolve()
    except ImportError:
        pass
    except Exception as exc:  # noqa: BLE001
        logger.warning("weasyprint PDF conversion failed: %s", exc)

    # 3. wkhtmltopdf
    if shutil.which("wkhtmltopdf"):
        try:
            html = markdown_to_html(markdown_text)
            html_tmp = out.with_suffix(".tmp.html")
            html_tmp.write_text(html, encoding="utf-8")
            subprocess.run(
                ["wkhtmltopdf", str(html_tmp), str(out)],
                capture_output=True,
                timeout=60,
                check=True,
            )
            html_tmp.unlink(missing_ok=True)
            logger.info("PDF written via wkhtmltopdf: %s", out)
            return out.resolve()
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            logger.warning("wkhtmltopdf failed: %s", exc)

    logger.warning(
        "No PDF converter available (pandoc, weasyprint, wkhtmltopdf). "
        "Install one to enable PDF export."
    )
    return None


def markdown_to_docx(
    markdown_text: str,
    output_path: Path | str,
) -> Path | None:
    """Convert Markdown to DOCX.

    Two-tier strategy (v2.7.10):
      1. **Pandoc** when available — gold standard. Preserves the full
         CommonMark feature set including nested tables, footnotes, and
         the heading hierarchy that the audit-report templates emit.
      2. **python-docx fallback** when pandoc is absent — bundled in the
         PyInstaller desktop install so end-users never see "DOCX export
         not available". Handles the audit-report subset: ``# / ## / ###``
         headings, ``**bold**`` runs, ``_italic_`` runs, fenced code
         blocks, ``-`` bullets, and ``---`` horizontal rules. Tables are
         rendered as plain paragraphs (the templates use them rarely).

    Args:
        markdown_text: Markdown source string.
        output_path: Destination ``.docx`` file path.

    Returns:
        Resolved Path of the written file on success, ``None`` if neither
        backend is available.
    """
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    # --- Tier 1: pandoc -----------------------------------------------
    if shutil.which("pandoc"):
        try:
            subprocess.run(
                [
                    "pandoc",
                    "--from=markdown",
                    "--to=docx",
                    f"--output={out}",
                ],
                input=markdown_text,
                capture_output=True,
                text=True,
                timeout=60,
                check=True,
            )
            logger.info("DOCX written via pandoc: %s", out)
            return out.resolve()
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            logger.warning("pandoc DOCX conversion failed (%s); falling back", exc)

    # --- Tier 2: python-docx ------------------------------------------
    try:
        return _markdown_to_docx_pythondocx(markdown_text, out)
    except ImportError:
        logger.warning(
            "Neither pandoc nor python-docx is available — DOCX export disabled. "
            "Install pandoc or `pip install python-docx`."
        )
        return None


def _markdown_to_docx_pythondocx(markdown_text: str, output_path: Path) -> Path | None:
    """python-docx Markdown→DOCX fallback for the desktop bundle.

    Recognises the audit-report Markdown dialect emitted by the Jinja2
    templates plus the per-finding sheets. NOT a full CommonMark parser
    — by design. Adding a full parser would balloon the dependency
    surface; this targeted converter is ~120 lines and ships with every
    PyInstaller build.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Document-wide style: 11pt body, dark-stone heading colour to match
    # the gemstone palette when viewed in Word's default light theme
    # (the body still reads cleanly on dark backgrounds in Reading Mode).
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    HEADING_COLOR = RGBColor(0x1F, 0x29, 0x37)  # near-black slate
    GOLD_ACCENT = RGBColor(0x8B, 0x69, 0x14)  # antique gold

    in_code_block = False
    code_buffer: list[str] = []

    def flush_code() -> None:
        if code_buffer:
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buffer))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            code_buffer.clear()

    def add_inline(paragraph, text: str) -> None:
        """Render bold/italic/code spans inside a paragraph.

        Tokenises on **…**, _…_, *…*, and `…` then emits a Run per span.
        Order matters — bold (**) before italic (* and _) so a literal
        `**word**` doesn't collapse to `*word*`.
        """
        # Pattern: ``code`` | **bold** | __bold__ | *italic* | _italic_
        token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)")
        for part in token_re.split(text):
            if not part:
                continue
            if part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            elif (part.startswith("**") and part.endswith("**")) or (
                part.startswith("__") and part.endswith("__")
            ):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif (part.startswith("*") and part.endswith("*")) or (
                part.startswith("_") and part.endswith("_")
            ):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()

        # Fenced code blocks
        if line.startswith("```"):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_buffer.append(line)
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
            continue

        # Headings
        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=0)
            for r in h.runs:
                r.font.color.rgb = GOLD_ACCENT
            continue
        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=1)
            for r in h.runs:
                r.font.color.rgb = HEADING_COLOR
            continue
        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=2)
            for r in h.runs:
                r.font.color.rgb = HEADING_COLOR
            continue

        # Bullet
        if line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:].strip())
            continue

        # Blank line
        if not line.strip():
            doc.add_paragraph()
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_inline(p, line)

    flush_code()
    doc.save(str(output_path))
    logger.info("DOCX written via python-docx fallback: %s", output_path)
    return output_path.resolve()


def get_available_formats() -> list[str]:
    """Return list of available output formats based on installed tools.

    Always includes ``"markdown"`` and ``"json"``.  Conditionally includes
    ``"html"``, ``"pdf"``, and ``"docx"`` based on what is installed.

    Returns:
        Sorted list of format name strings.
    """
    formats = ["json", "markdown"]

    # HTML is available if pandoc or the markdown library is present.
    has_pandoc = bool(shutil.which("pandoc"))
    has_md_lib = _can_import("markdown")
    if has_pandoc or has_md_lib:
        formats.append("html")

    # PDF requires pandoc, weasyprint, or wkhtmltopdf.
    has_weasyprint = _can_import("weasyprint")
    has_wkhtmltopdf = bool(shutil.which("wkhtmltopdf"))
    if has_pandoc or has_weasyprint or has_wkhtmltopdf:
        formats.append("pdf")

    # DOCX: pandoc preferred, python-docx fallback (v2.7.10 — bundled
    # in the desktop installer so DOCX export works without external
    # tooling).
    has_python_docx = _can_import("docx")
    if has_pandoc or has_python_docx:
        formats.append("docx")

    return sorted(formats)


# ---------------------------------------------------------------------------
# High-level export
# ---------------------------------------------------------------------------


def export_report(
    report: AuditReport,
    output_dir: Path | str,
    formats: list[str] | None = None,
    template_dir: Path | str = "templates",
    template_name: str = "audit_report.md",
) -> dict[str, Path]:
    """Export an AuditReport in multiple formats.

    Args:
        report: The AuditReport to export.
        output_dir: Directory to write files to (created if absent).
        formats: List of formats to generate.  Defaults to all available.
            Options: ``"json"``, ``"markdown"``, ``"html"``, ``"pdf"``,
            ``"docx"``.
        template_dir: Path to Jinja2 templates directory.
        template_name: Template filename to use for Markdown rendering.

    Returns:
        Dict mapping format name to the resolved output Path.
        e.g., ``{"json": Path("…/RPT-XYZ_report.json"), …}``
    """
    from oraculus_di_auditor.reporting.template_engine import ReportTemplateEngine

    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    available = get_available_formats()
    requested = {f.lower() for f in (formats or available)}
    stem = f"{report.report_id}_report"
    written: dict[str, Path] = {}

    if "json" in requested:
        written["json"] = _export_json(report, out_dir, stem)

    md_formats = requested & {"markdown", "html", "pdf", "docx"}
    if md_formats:
        engine = ReportTemplateEngine(template_dir=template_dir)
        md = engine.render_markdown(report, template_name=template_name)
        title = report.title or "ODIA Audit Report"
        written.update(_export_md_formats(md, md_formats, out_dir, stem, title))

    for fmt in requested - {"json", "markdown", "html", "pdf", "docx"}:
        logger.warning("Format '%s' not available — skipping.", fmt)

    return written


def _export_md_formats(
    md: str,
    formats: set[str],
    out_dir: Path,
    stem: str,
    title: str,
) -> dict[str, Path]:
    """Render all Markdown-derived formats from a pre-rendered Markdown string."""
    written: dict[str, Path] = {}
    if "markdown" in formats:
        written["markdown"] = _export_markdown(md, out_dir, stem)
    if "html" in formats:
        written["html"] = _export_html(md, out_dir, stem)
    if "pdf" in formats:
        result = _export_pdf(md, out_dir, stem, title)
        if result:
            written["pdf"] = result
    if "docx" in formats:
        result = _export_docx(md, out_dir, stem)
        if result:
            written["docx"] = result
    return written


def _export_json(report: AuditReport, out_dir: Path, stem: str) -> Path:
    """Write report as JSON and return the path."""
    path = out_dir / f"{stem}.json"
    _write_json(report, path)
    print(f"[OK] JSON  → {path}")
    return path


def _export_markdown(md: str, out_dir: Path, stem: str) -> Path:
    """Write rendered Markdown to disk and return the path."""
    path = out_dir / f"{stem}.md"
    path.write_text(md, encoding="utf-8")
    print(f"[OK] Markdown → {path}")
    return path


def _export_html(md: str, out_dir: Path, stem: str) -> Path:
    """Convert Markdown to HTML, write to disk, and return the path."""
    path = out_dir / f"{stem}.html"
    html = markdown_to_html(md)
    path.write_text(html, encoding="utf-8")
    print(f"[OK] HTML  → {path}")
    return path


def _export_pdf(md: str, out_dir: Path, stem: str, title: str) -> Path | None:
    """Convert Markdown to PDF and return the path, or None if unavailable."""
    path = out_dir / f"{stem}.pdf"
    result = markdown_to_pdf(md, path, title=title)
    if result:
        print(f"[OK] PDF   → {result}")
        return result
    print("[SKIP] PDF — no converter available")
    return None


def _export_docx(md: str, out_dir: Path, stem: str) -> Path | None:
    """Convert Markdown to DOCX and return the path, or None if unavailable."""
    path = out_dir / f"{stem}.docx"
    result = markdown_to_docx(md, path)
    if result:
        print(f"[OK] DOCX  → {result}")
        return result
    print("[SKIP] DOCX — pandoc not available")
    return None


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _can_import(module_name: str) -> bool:
    """Return True if *module_name* can be imported."""
    import importlib.util

    return importlib.util.find_spec(module_name) is not None


def _write_json(report: AuditReport, path: Path) -> None:
    """Serialise *report* to a pretty-printed JSON file."""
    data: dict[str, Any] = json.loads(report.model_dump_json())
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")

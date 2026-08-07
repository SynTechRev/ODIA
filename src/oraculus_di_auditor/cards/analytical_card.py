"""C.O.N.T.R.A. Analytical Card DOCX generator.

Produces a per-contract Analytical Card in the MAS document style:
Garamond typeface, malachite green (#1D6B44) section headings,
tan (#C8A882) sub-headings, page-of-total footer.

No em-dash or en-dash characters are written to the document (platform
convention; use hyphens or Unicode substitutes at the Python layer).

Source: C.O.N.T.R.A. Framework V1.0 Section VIII, Handoff Specification V1.0 Section 7.2
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

if TYPE_CHECKING:
    from ..contra.base import Finding
    from ..scoring.casi import CasiAxes

# Palette
_MALACHITE = RGBColor(0x1D, 0x6B, 0x44)  # #1D6B44 — section headings
_TAN = RGBColor(0xC8, 0xA8, 0x82)  # #C8A882 — sub-headings
_BLACK = RGBColor(0x00, 0x00, 0x00)

_FONT_BODY = "Garamond"
_FONT_FALLBACK = "Times New Roman"

# CASI axis display names
_AXIS_LABELS = {
    "remedy_foreclosure": "Remedy Foreclosure",
    "data_extraction_depth": "Data Extraction Depth",
    "modification_and_consent": "Modification and Consent",
    "procedural_adhesion": "Procedural Adhesion",
    "enforcement_cost_asymmetry": "Enforcement Cost Asymmetry",
}


@dataclass
class AnalyticalCardInput:
    """Input bundle for build_analytical_card()."""

    entity_name: str
    entity_id: str | None
    doc_type: str
    effective_date: str | None
    version_label: str | None
    document_hash: str
    source_url: str | None
    wayback_url: str | None
    findings: list[Finding] = field(default_factory=list)
    casi_axes: CasiAxes | None = None
    framework_version: str = "1.0"
    ingestion_date: str | None = None

    def __post_init__(self) -> None:
        if self.ingestion_date is None:
            self.ingestion_date = datetime.now(UTC).strftime("%Y-%m-%d")


def _set_run_font(run, size_pt: float = 11, bold: bool = False) -> None:
    run.font.name = _FONT_BODY
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # Fallback hint for Windows font substitution
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), _FONT_BODY)
    rFonts.set(qn("w:hAnsi"), _FONT_BODY)
    rFonts.set(qn("w:cs"), _FONT_FALLBACK)
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _add_section_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    _set_run_font(run, size_pt=14, bold=True)
    run.font.color.rgb = _MALACHITE
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)


def _add_sub_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, size_pt=11, bold=True)
    run.font.color.rgb = _TAN


def _add_body_para(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    for run in para.runs:
        _set_run_font(run)


def _add_kv(doc: Document, key: str, value: str) -> None:
    para = doc.add_paragraph()
    key_run = para.add_run(f"{key}: ")
    _set_run_font(key_run, bold=True)
    val_run = para.add_run(value)
    _set_run_font(val_run)


def _add_page_of_total_footer(doc: Document) -> None:
    """Add 'Page X of Y' footer using Word field codes."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    run = para.add_run("Page ")
    _set_run_font(run, size_pt=9)

    # PAGE field
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    run2 = para.add_run(" of ")
    _set_run_font(run2, size_pt=9)

    # NUMPAGES field
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "begin")
    run2._r.append(fld2)
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    run2._r.append(instr2)
    fld_end2 = OxmlElement("w:fldChar")
    fld_end2.set(qn("w:fldCharType"), "end")
    run2._r.append(fld_end2)


def _add_casi_table(doc: Document, axes: CasiAxes) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(["Axis", "Score (0-20)", "Contribution"]):
        run = hdr[i].paragraphs[0].add_run(text)
        _set_run_font(run, bold=True)
        run.font.color.rgb = _MALACHITE

    axes_dict = axes.to_dict()
    for axis_key, label in _AXIS_LABELS.items():
        score = axes_dict.get(axis_key, 0)
        pct = f"{score / 20 * 100:.0f}%"
        row = table.add_row().cells
        _set_run_font(row[0].paragraphs[0].add_run(label))
        _set_run_font(row[1].paragraphs[0].add_run(str(score)))
        _set_run_font(row[2].paragraphs[0].add_run(pct))

    # Aggregate row
    agg_row = table.add_row().cells
    agg_run = agg_row[0].paragraphs[0].add_run("AGGREGATE")
    _set_run_font(agg_run, bold=True)
    agg_score_run = agg_row[1].paragraphs[0].add_run(str(axes_dict["aggregate"]))
    _set_run_font(agg_score_run, bold=True)
    band_run = agg_row[2].paragraphs[0].add_run(axes_dict["band"])
    _set_run_font(band_run, bold=True)


def _add_findings_table(doc: Document, findings: list[Finding]) -> None:
    if not findings:
        _add_body_para(doc, "No findings for this document.")
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Layer/Sub", "Severity", "Anchor", "Excerpt (<=15 words)"]
    for i, h in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(h)
        _set_run_font(run, bold=True)
        run.font.color.rgb = _MALACHITE

    for f in findings:
        row = table.add_row().cells
        _set_run_font(row[0].paragraphs[0].add_run(f"{f.layer}.{f.sub_detector}"))
        sev_run = row[1].paragraphs[0].add_run(f.severity.value.upper())
        _set_run_font(sev_run, bold=(f.severity.value == "critical"))
        # Truncate anchor to fit cell
        short_anchor = (
            f.doctrinal_anchor[:60] + "..."
            if len(f.doctrinal_anchor) > 60
            else f.doctrinal_anchor
        )
        _set_run_font(row[2].paragraphs[0].add_run(short_anchor))
        _set_run_font(row[3].paragraphs[0].add_run(f.evidence_span.verbatim_excerpt))


def build_analytical_card(
    input_data: AnalyticalCardInput,
    output_dir: str | Path,
) -> str:
    """Generate a DOCX Analytical Card for one commercial document.

    Returns the absolute path to the generated .docx file.
    Filename: contra_card_{document_hash[:12]}.docx
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    safe_name = "".join(
        c if c.isalnum() or c in "-_" else "_" for c in input_data.entity_name[:40]
    )
    filename = f"contra_card_{safe_name}_{input_data.document_hash[:12]}.docx"
    output_path = output_dir / filename

    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("C.O.N.T.R.A. ANALYTICAL CARD")
    _set_run_font(title_run, size_pt=18, bold=True)
    title_run.font.color.rgb = _MALACHITE

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        f"Commercial Contract Asymmetry Analysis -- Framework V{input_data.framework_version}"
    )
    _set_run_font(sub_run, size_pt=11)
    sub_run.font.color.rgb = _TAN

    doc.add_paragraph()

    # Block 1: Identification
    _add_section_heading(doc, "I. IDENTIFICATION")
    _add_kv(doc, "Entity", input_data.entity_name)
    _add_kv(doc, "Document Type", input_data.doc_type)
    _add_kv(doc, "Effective Date", input_data.effective_date or "Not specified")
    _add_kv(doc, "Version", input_data.version_label or "Current")
    _add_kv(doc, "SHA-256", input_data.document_hash)
    _add_kv(doc, "Ingestion Date", input_data.ingestion_date or "")
    if input_data.source_url:
        _add_kv(doc, "Source URL", input_data.source_url)
    if input_data.wayback_url:
        _add_kv(doc, "Wayback Capture", input_data.wayback_url)

    # Block 2: CASI Score
    _add_section_heading(doc, "II. CONSUMER ADHESION SEVERITY INDEX (CASI)")
    if input_data.casi_axes is not None:
        axes_dict = input_data.casi_axes.to_dict()
        _add_kv(
            doc,
            "Aggregate",
            f"{axes_dict['aggregate']}/100 -- {axes_dict['band']}",
        )
        doc.add_paragraph()
        _add_casi_table(doc, input_data.casi_axes)
    else:
        _add_body_para(doc, "CASI scoring pending.")

    # Block 3: Detector Findings
    _add_section_heading(doc, "III. DETECTOR FINDINGS (L-11 through L-20)")
    finding_count = len(input_data.findings)
    critical_count = sum(
        1 for f in input_data.findings if f.severity.value == "critical"
    )
    _add_kv(doc, "Total Findings", str(finding_count))
    _add_kv(doc, "CRITICAL", str(critical_count))
    doc.add_paragraph()
    _add_findings_table(doc, input_data.findings)

    # Block 4: Recommended Actions (placeholder — populated by Phase G)
    _add_section_heading(doc, "IV. RECOMMENDED ACTIONS")
    _add_body_para(
        doc,
        "Recommended actions are populated during the full ingestion pipeline "
        "(Phase G). Channels assessed: CCPA requests, California Delete Act DROP, "
        "CPPA/FTC/AG regulatory complaint, small-claims candidates, "
        "CCP section 1281.97 default candidates, PAGA (if employment context).",
    )

    # Block 5: Chain of Custody
    _add_section_heading(doc, "V. CHAIN OF CUSTODY")
    _add_kv(doc, "Document Hash (SHA-256)", input_data.document_hash)
    _add_kv(
        doc,
        "Generated",
        datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC"),
    )

    _add_page_of_total_footer(doc)

    doc.save(str(output_path))
    return str(output_path)
